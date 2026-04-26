import importlib
import json
import os
import sys
import re
import uuid
import asyncio
import time
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional
from datetime import datetime
import sqlite3

# 将当前文件的上一级目录加入到搜索路径中
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)
# 🌟 引入多Agent工作流
from multi_agent_workflow.agent_workflow import multi_agent_app

import requests
import urllib3
from docx2pdf import convert

# 🌟 临时禁用 SSL 警告（仅用于本地开发环境）
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
from fastapi import FastAPI, HTTPException, BackgroundTasks, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from openai import OpenAI
from pdf2image import convert_from_path
from pydantic import BaseModel
from docxtpl import DocxTemplate, RichText
# 🌟 修改 1：去掉了废弃的表 ID 导入
from common.config import (
    FEISHU_APP_ID as APP_ID,
    FEISHU_APP_SECRET as APP_SECRET,
    FEISHU_APP_TOKEN as APP_TOKEN,
    FEISHU_TABLE_ID_JOBS as TABLE_ID,  # ✅ 统一使用总表
    FEISHU_TABLE_ID_PROMPTS,  # 🌟 新增：Prompt表
    FEISHU_TABLE_ID_RESUMES,  # 🌟 新增：简历表
    LLM_API_KEY,
    LLM_BASE_URL,
    LLM_MODEL,
    VISION_LLM_MODEL,
    SERPER_API_KEY,
)
import common.config as _config_module


app = FastAPI(title="简历生成与飞书同步后台")

# 允许本地前端跨域访问
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000", "http://127.0.0.1:3000",
        "http://localhost:3001", "http://127.0.0.1:3001",
        "http://localhost:3002", "http://127.0.0.1:3002",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = (BASE_DIR.parent / "boss-cli" / "自动改写简历模板FOR GEMINI.docx").resolve()
EXPORT_TEMPLATE_PATH = BASE_DIR / "template.docx"
TEMPLATES_DIR = BASE_DIR / "templates"
DEFAULT_TEMPLATE_MARKER = TEMPLATES_DIR / ".default"
ACTIVE_TEMPLATE_FILE = TEMPLATES_DIR / "active_template.txt"
TEMP_DOCX_PATH = BASE_DIR / "temp_resume.docx"
TEMP_PDF_PATH = BASE_DIR / "temp_resume.pdf"
SETTINGS_DATA_PATH = BASE_DIR / "data" / "settings.json"
DB_PATH = BASE_DIR.parent / "data" / "job_hunter.db"
client = OpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL) if LLM_API_KEY else None

# 导入飞书API模块
_AI_AGENTS_DIR = os.path.join(PROJECT_ROOT, "ai_agents")
if _AI_AGENTS_DIR not in sys.path:
    sys.path.insert(0, _AI_AGENTS_DIR)
from common import feishu_api
from qa_evaluator import qa_evaluate_resume

# 🌟 导入真实的 AI 评估和改写函数
from ai_evaluator import evaluate_single_job, load_resume as load_resume_evaluator
from ai_scorer import get_job_match_score, rewrite_resume_for_job, generate_greeting, deep_evaluate_resume, parse_resume_markdown

# 全局任务队列管理（用于 SSE 实时日志推送）
task_queues: Dict[str, asyncio.Queue] = {}
task_status: Dict[str, Dict[str, Any]] = {}

# 🌟 全局后台任务状态字典，用于短轮询
GLOBAL_TASK_STATE = {"is_processing": False}


def parse_to_richtext(text: Any, doc: Optional[DocxTemplate] = None) -> RichText:
    """
    将 Markdown 风格的 **粗体** 转为 Word RichText 对象，并保留换行。
    新增：自动识别 URL 并转换为真实可点击的超链接。
    """
    if text is None or (isinstance(text, str) and not text.strip()):
        rt = RichText("")
        return rt
    lines = str(text).split("\n")
    rt = RichText("")
    for i, line in enumerate(lines):
        if line.count(" · ") == 2:
            line = line.replace(" · ", "\t")

        bold_parts = re.split(r"(\*\*.*?\*\*)", line)
        for bold_part in bold_parts:
            if bold_part.startswith("**") and bold_part.endswith("**"):
                inner_text = bold_part[2:-2]
                url_parts = re.split(r"(https?://[^\s]+)", inner_text)
                for url_part in url_parts:
                    if re.match(r"https?://[^\s]+", url_part):
                        if doc:
                            url_id = doc.build_url_id(url_part)
                            rt.add(url_part, bold=True, font="微软雅黑", size=18, color="#0563C1", underline=True, url_id=url_id)
                        else:
                            rt.add(url_part, bold=True, font="微软雅黑", size=18, color="#0563C1", underline=True)
                    elif url_part:
                        rt.add(url_part, bold=True, font="微软雅黑", size=18)
            elif bold_part:
                url_parts = re.split(r"(https?://[^\s]+)", bold_part)
                for url_part in url_parts:
                    if re.match(r"https?://[^\s]+", url_part):
                        if doc:
                            url_id = doc.build_url_id(url_part)
                            rt.add(url_part, font="微软雅黑", size=18, color="#0563C1", underline=True, url_id=url_id)
                        else:
                            rt.add(url_part, font="微软雅黑", size=18, color="#0563C1", underline=True)
                    elif url_part:
                        rt.add(url_part, font="微软雅黑", size=18)

        if i < len(lines) - 1:
            rt.add("\a")  

    return rt


def sanitize_filename_component(raw: Any, default: str) -> str:
    if raw is None:
        return default
    s = str(raw).strip()
    if not s:
        return default
    for ch in '\\/:*?"<>|\n\r\t\x00':
        s = s.replace(ch, "")
    s = re.sub(r"\s+", " ", s).strip()
    s = s.strip(" .")
    if not s:
        return default
    return s


def feishu_field_to_plain_str(val: Any, default: str = "") -> str:
    if val is None:
        return default
    if isinstance(val, str):
        return val or default
    if isinstance(val, (int, float)):
        return str(val)
    if isinstance(val, list):
        chunks: List[str] = []
        for item in val:
            if isinstance(item, str):
                chunks.append(item)
            elif isinstance(item, dict):
                text = item.get("text") or item.get("name") or item.get("value")
                if text is not None:
                    chunks.append(str(text))
            elif item is not None:
                chunks.append(str(item))
        return "".join([c for c in chunks if c]) or default
    if isinstance(val, dict):
        text = val.get("text") or val.get("name") or val.get("value")
        if text is not None:
            return str(text)
        return str(val) or default
    return str(val) or default


def fetch_bitable_record_by_id(
    tenant_access_token: str,
    app_token: str,
    table_id: str,
    record_id: str,
) -> Optional[Dict[str, Any]]:
    url = (
        f"https://open.feishu.cn/open-apis/bitable/v1/apps/{app_token}"
        f"/tables/{table_id}/records/{record_id}"
    )
    headers = {"Authorization": f"Bearer {tenant_access_token}"}
    try:
        resp = requests.get(url, headers=headers, timeout=15, proxies={"http": None, "https": None}, verify=False)
        resp.raise_for_status()
        data = resp.json()
    except requests.exceptions.SSLError as ssl_err:
        print(f"❌ SSL 错误 (fetch_bitable_record_by_id): {ssl_err}")
        return None
    except requests.exceptions.RequestException as req_err:
        print(f"❌ 网络请求错误 (fetch_bitable_record_by_id): {req_err}")
        return None
    except ValueError as val_err:
        print(f"❌ JSON 解析错误 (fetch_bitable_record_by_id): {val_err}")
        return None
    if data.get("code") != 0:
        return None
    return data.get("data", {}).get("record")


def _require_env(name: str, value: Optional[str]) -> str:
    if not value:
        raise HTTPException(status_code=500, detail=f"缺少环境变量: {name}")
    return value

# 🌟 新增：main 专属的 Token 缓存池
_MAIN_TOKEN_CACHE = {"token": None, "expires_at": 0}

def get_tenant_access_token(app_id: str, app_secret: str) -> str:
    global _MAIN_TOKEN_CACHE
    current_time = time.time()
    
    # 🌟 命中缓存，直接秒回
    if _MAIN_TOKEN_CACHE["token"] and current_time < _MAIN_TOKEN_CACHE["expires_at"] - 300:
        return _MAIN_TOKEN_CACHE["token"]

    url = "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal"
    try:
        resp = requests.post(
            url,
            json={"app_id": app_id, "app_secret": app_secret},
            timeout=15,
            proxies={"http": None, "https": None},
            verify=False
        )
        resp.raise_for_status()
        data = resp.json()
    except requests.exceptions.SSLError as ssl_err:
        print(f"❌ SSL 错误 (get_tenant_access_token): {ssl_err}")
        raise HTTPException(
            status_code=502,
            detail=f"飞书鉴权 SSL 错误，请检查网络代理设置: {str(ssl_err)}",
        )
    except requests.exceptions.RequestException as req_err:
        print(f"❌ 网络请求错误 (get_tenant_access_token): {req_err}")
        raise HTTPException(
            status_code=502,
            detail=f"飞书鉴权网络错误: {str(req_err)}",
        )

    if data.get("code") != 0:
        raise HTTPException(
            status_code=502,
            detail=f"飞书鉴权失败: {data.get('msg', 'unknown error')}",
        )
        
    # 存入缓存
    _MAIN_TOKEN_CACHE["token"] = data["tenant_access_token"]
    _MAIN_TOKEN_CACHE["expires_at"] = current_time + data.get("expire", 7200)
    return _MAIN_TOKEN_CACHE["token"]


def fetch_bitable_records(
    tenant_access_token: str,
    app_token: str,
    table_id: str,
) -> List[Dict[str, Any]]:
    url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{app_token}/tables/{table_id}/records"
    headers = {"Authorization": f"Bearer {tenant_access_token}"}

    all_items: List[Dict[str, Any]] = []
    page_token: Optional[str] = None

    while True:
        params: Dict[str, Any] = {"page_size": 100}
        if page_token:
            params["page_token"] = page_token

        _max_retries = 3
        _last_err: Optional[Exception] = None
        data = None
        for _attempt in range(_max_retries):
            try:
                resp = requests.get(url, headers=headers, params=params, timeout=15, proxies={"http": None, "https": None}, verify=False)
                resp.raise_for_status()
                data = resp.json()
                _last_err = None
                break
            except (requests.exceptions.SSLError, requests.exceptions.ConnectionError) as net_err:
                _last_err = net_err
                print(f"⚠️ 网络/SSL 错误 (第 {_attempt+1}/{_max_retries} 次), 稍后重试: {net_err}")
                import time as _time; _time.sleep(1.5 * (_attempt + 1))
            except requests.exceptions.RequestException as req_err:
                _last_err = req_err
                break

        if _last_err is not None:
            if isinstance(_last_err, (requests.exceptions.SSLError, requests.exceptions.ConnectionError)):
                print(f"❌ SSL/网络错误已达最大重试次数 (fetch_bitable_records): {_last_err}")
                raise HTTPException(
                    status_code=502,
                    detail=f"读取飞书表格 SSL 错误（已重试 {_max_retries} 次）: {str(_last_err)}",
                )
            else:
                print(f"❌ 网络请求错误 (fetch_bitable_records): {_last_err}")
                raise HTTPException(
                    status_code=502,
                    detail=f"读取飞书表格网络错误: {str(_last_err)}",
                )

        if data.get("code") != 0:
            raise HTTPException(
                status_code=502,
                detail=f"读取飞书表格失败: {data.get('msg', 'unknown error')}",
            )

        payload = data.get("data", {})
        items = payload.get("items", [])
        all_items.extend(items)

        if not payload.get("has_more"):
            break
        page_token = payload.get("page_token")
        if not page_token:
            break

    return all_items

def get_active_resume_from_feishu() -> str:
    """从飞书配置中心读取【启用】状态的简历内容"""
    try:
        token = get_tenant_access_token(APP_ID, APP_SECRET)
        url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{FEISHU_TABLE_ID_RESUMES}/records/search"
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
        payload = {
            "filter": {
                "conjunction": "and",
                "conditions": [{"field_name": "当前状态", "operator": "is", "value": ["启用"]}]
            }
        }
        resp = requests.post(url, headers=headers, json=payload, timeout=15, proxies={"http": None, "https": None}, verify=False)
        resp.raise_for_status()
        items = resp.json().get("data", {}).get("items", [])
        if items:
            return feishu_field_to_plain_str(items[0].get("fields", {}).get("简历内容", ""))
        return ""
    except requests.exceptions.SSLError as ssl_err:
        print(f"❌ SSL 错误 (get_active_resume_from_feishu): {ssl_err}")
        return ""
    except requests.exceptions.RequestException as req_err:
        print(f"❌ 网络请求错误 (get_active_resume_from_feishu): {req_err}")
        return ""
    except Exception as e:
        print(f"❌ 读取飞书云端简历失败: {e}")
        return ""

def normalize_job_record(record: Dict[str, Any], default_platform: str = "未知") -> Optional[Dict[str, Any]]:
    fields = record.get("fields", {})

    def safe_get_int(key, default=0):
        val = fields.get(key)
        try:
            return int(val) if val is not None else default
        except (ValueError, TypeError):
            return default

    def safe_get_text(key: str, default: str = "") -> str:
        val = fields.get(key)
        if val is None:
            return default
        if isinstance(val, str):
            return val
        if isinstance(val, (int, float)):
            return str(val)
        if isinstance(val, list):
            chunks: List[str] = []
            for item in val:
                if isinstance(item, str):
                    chunks.append(item)
                elif isinstance(item, dict):
                    text = item.get("text") or item.get("name") or item.get("value")
                    if text is not None:
                        chunks.append(str(text))
                elif item is not None:
                    chunks.append(str(item))
            return "".join([c for c in chunks if c]) or default
        if isinstance(val, dict):
            text = val.get("text") or val.get("name") or val.get("value")
            if text is not None:
                return str(text)
            return str(val)
        return str(val)

    def safe_get_link(key: str) -> str:
        val = fields.get(key)
        if val is None:
            return ""
        if isinstance(val, str):
            return val.strip()
        if isinstance(val, dict):
            link = val.get("link") or val.get("text") or ""
            return str(link).strip()
        return ""

    def safe_get_list(key: str) -> List[str]:
        val = fields.get(key)
        if val is None:
            return []
        if isinstance(val, list):
            result: List[str] = []
            for item in val:
                if isinstance(item, str):
                    if item.strip():
                        result.append(item.strip())
                elif isinstance(item, dict):
                    text = item.get("text") or item.get("name") or item.get("value")
                    if text:
                        result.append(str(text).strip())
                elif item is not None:
                    result.append(str(item).strip())
            return [x for x in result if x]
        text = safe_get_text(key, "")
        if not text:
            return []
        return [part.strip() for part in str(text).replace("；", ",").replace("，", ",").split(",") if part.strip()]

    return {
        "record_id": record.get("record_id"),
        "job_name": fields.get("岗位名称", "未知岗位"),
        "company_name": fields.get("公司名称", "未知公司"),
        "city": fields.get("城市", "未知城市"),
        "salary": fields.get("薪资", "面议"),
        "follow_status": fields.get("跟进状态", "新线索"),
        "scale": fields.get("公司规模", "规模不详"),
        "industry": fields.get("所属行业", "未知行业"),
        "education": fields.get("学历要求", "学历不限"),
        "experience": fields.get("经验要求", "经验不限"),
        "ai_score": safe_get_int("AI总分", 0),
        "bg_score": safe_get_int("背景得分", 0),
        "skill_score": safe_get_int("技能得分", 0),
        "exp_score": safe_get_int("经验得分", 0),
        "job_detail": fields.get("岗位详情", "暂无详情"),
        "skill_req": safe_get_text("技能要求", ""),
        "hr_skills": safe_get_list("HR技能标签"),
        "benefits": safe_get_list("福利标签"),
        "hr_active": safe_get_text("HR活跃度", ""),
        "delivery_date": safe_get_text("投递日期", ""),
        "fetch_time": safe_get_text("抓取时间", ""),
        "work_address": safe_get_text("工作地址", ""),
        "my_review": safe_get_text("我的复核", ""),
        "ai_rewrite_json": safe_get_text("AI改写JSON", "") or safe_get_text("AI改写简历", ""),
        "manual_refined_resume": safe_get_text("人工精修版简历", ""),
        "dream_picture": safe_get_text("理想画像与能力信号", ""),
        "ats_ability_analysis": safe_get_text("核心能力词典", ""),
        "strong_fit_assessment": safe_get_text("高杠杆匹配点", ""),
        "risk_red_flags": safe_get_text("致命硬伤与毒点", ""),
        "deep_action_plan": safe_get_text("破局行动计划", ""),
        "greeting_msg": safe_get_text("打招呼语", ""),
        "platform": safe_get_text("招聘平台", "") or safe_get_text("数据来源", "") or default_platform,
        "role": safe_get_text("角色", "") or safe_get_text("发布人角色", "") or "未知",
        "publish_date": safe_get_text("发布日期", ""),
        "preliminary_score": safe_get_int("初步打分", 0),
        "bonus_words": safe_get_text("加分词", ""),
        "deduction_words": safe_get_text("减分词", ""),
        "grade": safe_get_text("综合评级 (A-F)", "") or safe_get_text("综合评级", ""),
        "ai_evaluation_detail": safe_get_text("AI评估详情", ""),
        "role_match": safe_get_int("核心-角色匹配", 0),
        "skills_align": safe_get_int("核心-技能重合", 0),
        "seniority": safe_get_int("高权-职级资历", 0),
        "compensation": safe_get_int("高权-薪资契合", 0),
        "interview_prob": safe_get_int("高权-面试概率", 0),
        "work_mode": safe_get_int("中权-工作模式", 0),
        "company_stage": safe_get_int("中权-公司阶段", 0),
        "market_fit": safe_get_int("中权-赛道前景", 0),
        "growth": safe_get_int("中权-成长空间", 0),
        "timeline": safe_get_int("低权-招聘周期", 0),
        "job_link": safe_get_link("岗位链接"),
    }


def upload_file_to_feishu(
    tenant_access_token: str,
    app_token: str,
    file_path: Path,
    file_name: Optional[str] = None,
) -> str:
    url = "https://open.feishu.cn/open-apis/drive/v1/medias/upload_all"
    headers = {"Authorization": f"Bearer {tenant_access_token}"}

    display_name = file_name if file_name else file_path.name

    content_type = "application/octet-stream"
    if file_path.suffix.lower() == ".docx":
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif file_path.suffix.lower() == ".pdf":
        content_type = "application/pdf"
    elif file_path.suffix.lower() in {".jpg", ".jpeg"}:
        content_type = "image/jpeg"

    data = {
        "file_name": display_name,
        "parent_type": "bitable_file",
        "parent_node": app_token,
        "size": str(file_path.stat().st_size),
    }

    try:
        with file_path.open("rb") as f:
            files = {"file": (display_name, f, content_type)}
            resp = requests.post(url, headers=headers, data=data, files=files, timeout=60, proxies={"http": None, "https": None}, verify=False)

        resp.raise_for_status()
        result = resp.json()
    except requests.exceptions.SSLError as ssl_err:
        print(f"❌ SSL 错误 (upload_file_to_feishu): {ssl_err}")
        raise HTTPException(
            status_code=502,
            detail=f"上传文件到飞书 SSL 错误({display_name})，请检查网络代理设置: {str(ssl_err)}",
        )
    except requests.exceptions.RequestException as req_err:
        print(f"❌ 网络请求错误 (upload_file_to_feishu): {req_err}")
        raise HTTPException(
            status_code=502,
            detail=f"上传文件到飞书网络错误({display_name}): {str(req_err)}",
        )
    
    if result.get("code") != 0:
        raise HTTPException(
            status_code=502,
            detail=f"上传文件到飞书失败({display_name}): {result.get('msg', 'unknown error')}",
        )
    return result["data"]["file_token"]


def update_bitable_record_attachments(
    tenant_access_token: str,
    app_token: str,
    table_id: str,
    record_id: str,
    word_token: str,
    pdf_token: str,
    image_tokens: List[str],
) -> None:
    url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{app_token}/tables/{table_id}/records/{record_id}"
    headers = {"Authorization": f"Bearer {tenant_access_token}"}
    payload = {
        "fields": {
            "Word附件": [{"file_token": word_token}],
            "PDF备份": [{"file_token": pdf_token}],
            "图片保存": [{"file_token": token} for token in image_tokens],
        }
    }
    try:
        resp = requests.put(url, headers=headers, json=payload, timeout=15, proxies={"http": None, "https": None}, verify=False)
        resp.raise_for_status()
        data = resp.json()
    except requests.exceptions.SSLError as ssl_err:
        print(f"❌ SSL 错误 (update_bitable_record_attachments): {ssl_err}")
        raise HTTPException(
            status_code=502,
            detail=f"回写飞书记录 SSL 错误，请检查网络代理设置: {str(ssl_err)}",
        )
    except requests.exceptions.RequestException as req_err:
        print(f"❌ 网络请求错误 (update_bitable_record_attachments): {req_err}")
        raise HTTPException(
            status_code=502,
            detail=f"回写飞书记录网络错误: {str(req_err)}",
        )
    
    if data.get("code") != 0:
        raise HTTPException(
            status_code=502,
            detail=f"回写飞书记录失败: {data.get('msg', 'unknown error')}",
        )


class SaveResumeRequest(BaseModel):
    record_id: str
    company_name: str
    target_job: str
    resume_data: Dict[str, Any]


class AIPolishRequest(BaseModel):
    selected_text: str
    instruction: str


class ExportRequest(BaseModel):
    job_id: str
    export_type: str
    resume_data: Dict[str, Any]
    platform: Optional[str] = None
    template_name: Optional[str] = None


class UpdateJobStatusRequest(BaseModel):
    job_id: str
    status: str
    platform: str = "BOSS直聘" 


class UpdateReviewCommentsRequest(BaseModel):
    job_id: str
    comments: str


class UpdateGreetingRequest(BaseModel):
    job_id: str
    greeting: str


class BatchDeleteRequest(BaseModel):
    job_ids: List[str]


class QAEvaluateRequest(BaseModel):
    job_id: str
    job_description: str
    resume_text: str
    platform: str = "BOSS直聘"


class SaveManualResumeRequest(BaseModel):
    job_id: str
    resume_text: str
    platform: str = "BOSS直聘"


class CopilotChatRequest(BaseModel):
    user_question: str
    context: Dict[str, Any]
    history: List[Dict[str, str]]


class BatchTaskRequest(BaseModel):
    task_type: str
    job_ids: List[str]
    scheduled_at: Optional[str] = None

class SaveConfigRequest(BaseModel):
    table_type: str # "resume" 或 "prompt"
    record_id: Optional[str] = None
    fields: Dict[str, Any]

class ToggleResumeStatusRequest(BaseModel):
    record_id: str

class ActivateResumeRequest(BaseModel):
    record_id: str

class MultiAgentRewriteRequest(BaseModel):
    job_id: str             # 岗位ID (用于拉取JD)
    original_resume: str    # 需要改写的原始经历片段或全文
    jd_text: str            # 岗位详情

class SettingsPayload(BaseModel):
    OPENAI_API_KEY: Optional[str] = None
    OPENAI_BASE_URL: Optional[str] = None
    SERPER_API_KEY: Optional[str] = None
    FEISHU_APP_ID: Optional[str] = None
    FEISHU_APP_SECRET: Optional[str] = None
    FEISHU_APP_TOKEN: Optional[str] = None


def extract_record_id(job_id: str) -> str:
    if not job_id:
        return job_id
    if "-" in job_id:
        parts = job_id.split("-")
        for part in reversed(parts):
            if part.startswith("rec"):
                return part
        return parts[-1]
    return job_id

# ==================== 🌟 策略中心(Strategy Lab) 专属 API ====================

@app.get("/api/strategy/config")
def get_strategy_config():
    """一键获取所有的简历和Prompt配置，供前端大盘渲染"""
    try:
        token = get_tenant_access_token(APP_ID, APP_SECRET)
        resumes = fetch_bitable_records(token, APP_TOKEN, FEISHU_TABLE_ID_RESUMES)
        prompts = fetch_bitable_records(token, APP_TOKEN, FEISHU_TABLE_ID_PROMPTS)
        
        # 清洗格式给前端
        resume_list = [{
            "record_id": r.get("record_id"),
            "version_name": feishu_field_to_plain_str(r.get("fields", {}).get("简历版本", "")),
            "content": feishu_field_to_plain_str(r.get("fields", {}).get("简历内容", "")),
            "status": feishu_field_to_plain_str(r.get("fields", {}).get("当前状态", "停用"))
        } for r in resumes]
        
        prompt_list = [{
            "record_id": r.get("record_id"),
            "strategy_name": feishu_field_to_plain_str(r.get("fields", {}).get("策略名称", "")),
            "content": feishu_field_to_plain_str(r.get("fields", {}).get("Prompt内容", "")),
            "status": feishu_field_to_plain_str(r.get("fields", {}).get("当前状态", "停用"))
        } for r in prompts]
        
        return {"status": "success", "resumes": resume_list, "prompts": prompt_list}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取策略配置失败: {e}")

@app.post("/api/strategy/toggle_resume_status")
def toggle_resume_status(payload: ToggleResumeStatusRequest):
    """将目标简历设为「启用」，同时将其余所有简历设为「已停用」（唯一启用排斥逻辑）"""
    try:
        token = get_tenant_access_token(APP_ID, APP_SECRET)
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

        # 1. 获取该表所有简历记录
        all_resumes = fetch_bitable_records(token, APP_TOKEN, FEISHU_TABLE_ID_RESUMES)

        # 2. 将其他「启用」/「启用中」的简历全部改为「已停用」
        for r in all_resumes:
            rid = r.get("record_id")
            if not rid or rid == payload.record_id:
                continue
            current_status = feishu_field_to_plain_str(r.get("fields", {}).get("当前状态", ""))
            if current_status in ("启用", "启用中"):
                deactivate_url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{FEISHU_TABLE_ID_RESUMES}/records/{rid}"
                resp = requests.put(deactivate_url, headers=headers, json={"fields": {"当前状态": "已停用"}}, timeout=15, proxies={"http": None, "https": None}, verify=False)
                resp.raise_for_status()
                if resp.json().get("code") != 0:
                    print(f"⚠️ 停用简历 {rid} 失败: {resp.json().get('msg')}")

        # 3. 将目标简历设为「启用」
        activate_url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{FEISHU_TABLE_ID_RESUMES}/records/{payload.record_id}"
        resp = requests.put(activate_url, headers=headers, json={"fields": {"当前状态": "启用"}}, timeout=15, proxies={"http": None, "https": None}, verify=False)
        resp.raise_for_status()
        data = resp.json()
        if data.get("code") != 0:
            raise Exception(f"更新目标简历状态失败: {data.get('msg')}")

        return {"status": "success", "message": f"已将 {payload.record_id} 设为启用，其余简历已停用"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"切换简历状态失败: {str(e)}")

@app.post("/api/strategy/activate_resume")
def activate_resume(payload: ActivateResumeRequest):
    """唯一启用排斥（精准版）：仅检索当前已启用的记录再批量停用，避免全量拉取"""
    try:
        token = get_tenant_access_token(APP_ID, APP_SECRET)
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

        # 1. 用 filter 只搜索「启用」/「启用中」的记录（比全量拉取更高效）
        search_url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{FEISHU_TABLE_ID_RESUMES}/records/search"
        search_payload = {
            "filter": {
                "conjunction": "or",
                "conditions": [
                    {"field_name": "当前状态", "operator": "is", "value": ["启用"]},
                    {"field_name": "当前状态", "operator": "is", "value": ["启用中"]},
                ]
            }
        }
        resp = requests.post(search_url, headers=headers, json=search_payload, timeout=15, proxies={"http": None, "https": None}, verify=False)
        resp.raise_for_status()
        active_items = resp.json().get("data", {}).get("items", [])

        # 2. 将其他已启用简历全部停用
        for r in active_items:
            rid = r.get("record_id")
            if not rid or rid == payload.record_id:
                continue
            deactivate_url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{FEISHU_TABLE_ID_RESUMES}/records/{rid}"
            resp_d = requests.put(deactivate_url, headers=headers, json={"fields": {"当前状态": "停用"}}, timeout=15, proxies={"http": None, "https": None}, verify=False)
            resp_d.raise_for_status()
            if resp_d.json().get("code") != 0:
                print(f"⚠️ 停用简历 {rid} 失败: {resp_d.json().get('msg')}")

        # 3. 启用目标简历
        activate_url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{FEISHU_TABLE_ID_RESUMES}/records/{payload.record_id}"
        resp_a = requests.put(activate_url, headers=headers, json={"fields": {"当前状态": "启用"}}, timeout=15, proxies={"http": None, "https": None}, verify=False)
        resp_a.raise_for_status()
        result = resp_a.json()
        if result.get("code") != 0:
            raise Exception(f"启用目标简历失败: {result.get('msg')}")

        return {"status": "success", "message": "已切换生效底稿"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"切换简历状态失败: {str(e)}")

# ==================== 🌟 新增：沙盒测试 (Playground) API ====================

class PromptTestRequest(BaseModel):
    prompt_content: str
    jd_text: str
    test_mode: str  # 新增：告诉后端当前测试的是什么类型 ("evaluate", "rewrite", "greeting")

# ==================== 🌟 沉浸工作台 (Immersive Workspace) API ====================

# ==================== 🌟 沉浸工作台 (Immersive Workspace) API ====================

@app.get("/api/jobs")
@app.get("/jobs")
async def get_all_jobs():
    """获取所有岗位列表，供沉浸工作台渲染使用"""
    print("\n====== 📡 沉浸工作台正在请求拉取岗位列表 ======")
    try:
        # 1. 获取飞书 Token
        token = get_tenant_access_token(APP_ID, APP_SECRET)
        
        # 2. 从飞书拉取总表所有数据
        raw_records = await asyncio.to_thread(
            fetch_bitable_records, token, APP_TOKEN, TABLE_ID
        )
        
        # 3. 数据清洗
        jobs = []
        for r in raw_records:
            normalized = normalize_job_record(r)
            if normalized:
                jobs.append(normalized)
                
        print(f"✅ 成功从飞书拉取并清洗了 {len(jobs)} 个岗位数据！")
        
       # 返回带有 items 外壳的对象，完美适配前端的解析逻辑
        return {"items": jobs, "status": "success"}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ 获取岗位列表失败: {e}")
        raise HTTPException(status_code=500, detail=f"获取岗位列表失败: {str(e)}")

# ==================== 🌟 极速录入 (Quick Import) API ====================

class JobImportTextRequest(BaseModel):
    raw_text: str

@app.post("/api/jobs/import/text")
async def import_job_from_text(payload: JobImportTextRequest):
    import json
    import re
    print("\n====== 📡 收到【极速录入】请求 ======")
    print(f"--> 收到文本长度: {len(payload.raw_text)} 字符")

    # ── Step 1: 正则预提取 URL 并判定平台 ─────────────────────────────────
    _url_match = re.search(r'https?://[^\s\u4e00-\u9fff，。！？\)\]]+', payload.raw_text)
    detected_url = _url_match.group(0).rstrip('.,;') if _url_match else ""

    def _detect_platform(url: str) -> str:
        if "liepin.com" in url:
            return "猎聘-极速录入"
        if "zhipin.com" in url:
            return "boss直聘-极速录入"
        if "51job.com" in url:
            return "51job-极速录入"
        if "zhaopin.com" in url:
            return "智联-极速录入"
        return "未知-极速录入"

    detected_platform = _detect_platform(detected_url)
    print(f"--> 预检测 URL: {detected_url or '(无)'} | 平台: {detected_platform}")

    system_prompt = """你是一个极度专业的招聘信息解析引擎。
从用户提供的招聘文本中精准提取字段，并【必须】以合法 JSON 对象输出，不得包含任何 Markdown 代码块标记或多余解释。

【严格提取规则】：
1. 输出必须是且只能是一个合法 JSON 对象，绝不允许任何额外文字。
2. 必须包含以下所有 key：job_name, company_name, city, salary, experience, education, skill_req, job_detail, publish_date, job_link。
3. 【猎头岗位特殊规则】：如果招聘文本由猎头（Headhunter）或猎头公司发布，雇主公司名称未明确写出，则 company_name 必须填写猎头公司名称，并在名称前加上"[猎头]"标识，例如："[猎头] 瑞贝斯人力资源"。绝不允许 company_name 为空或"未知"！
4. 如果某字段在文本中实在无法找到，填写 "-"，不允许留空。
5. skill_req 字段：提取所有技能要求，拼接为一行字符串，用逗号分隔。
6. job_detail 字段：尽可能完整地保留职位职责、任职要求等所有正文信息。
7. publish_date 字段：如果文本包含"今日更新"或"今日"，请在此字段原样输出"今日更新"（后端会转换为真实日期）；否则提取实际日期字符串，找不到则填"-"。
8. job_link 字段：如果文本中能找到以 http 开头的 URL，请将其完整放入此字段；找不到则填"-"。

输出示例：
{"job_name": "高级产品经理", "company_name": "[猎头] 猎途人力", "city": "上海", "salary": "30-45K·15薪", "experience": "5年以上", "education": "本科及以上", "skill_req": "产品规划,数据分析,用户研究", "job_detail": "负责...", "publish_date": "今日更新", "job_link": "https://..."}"""  # noqa

    def process_llm_and_feishu(text):
        print(">>> [1/3] 正在建立与大模型的专属连接...")
        from openai import OpenAI
        from common import feishu_api
        from datetime import datetime

        temp_client = OpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL)

        print(">>> [2/3] 正在发送给大模型提取信息... (请耐心等待)")
        try:
            response = temp_client.chat.completions.create(
                model=LLM_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"【待解析招聘文本】：\n{text}"}
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )
        except Exception:
            response = temp_client.chat.completions.create(
                model=LLM_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"【待解析招聘文本】：\n{text}"}
                ],
                temperature=0.1
            )

        raw = (response.choices[0].message.content or "").strip()
        print(f">>> ✅ 大模型成功返回数据！原始长度 {len(raw)} 字符")
        print(f"--- 原始返回前300字 ---\n{raw[:300]}")

        if raw.startswith("```json"):
            raw = raw[7:]
        elif raw.startswith("```"):
            raw = raw[3:]
        if raw.endswith("```"):
            raw = raw[:-3]
        raw = raw.strip()

        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError as e:
            print(f"❌ JSON 解析失败！原始内容:\n{raw}")
            raise ValueError(f"大模型返回的内容无法解析为 JSON: {e}")

        # ── 字段安全取値 + 兑底 ──
        company_name = parsed.get("company_name") or "未知猎头/公司"
        job_name     = parsed.get("job_name") or "未解析出岗位名"
        city         = parsed.get("city") or "-"
        salary       = parsed.get("salary") or "-"
        experience   = parsed.get("experience") or "-"
        education    = parsed.get("education") or "-"
        skill_req    = parsed.get("skill_req") or "-"
        job_detail   = parsed.get("job_detail") or "-"

        # ── 日期智能化：“今日更新” 转化为真实日期 ──
        raw_date = parsed.get("publish_date") or "-"
        publish_date = datetime.now().strftime("%Y-%m-%d") if "今日" in str(raw_date) else raw_date

        # ── 岗位链接：大模型提取为主，正则保底；空值或非 http 则不传给飞书 ──
        llm_link = parsed.get("job_link") or ""
        job_link = llm_link if (llm_link and llm_link != "-") else detected_url
        valid_url = job_link if (job_link and job_link.startswith("http")) else None
        print(f"--> 最终岗位链接: {valid_url or '(不写入，避免 URLFieldConvFail)'}")

        feishu_data = {
            "公司名称": str(company_name),
            "岗位名称": str(job_name),
            "城市": str(city),
            "薪资": str(salary),
            "经验要求": str(experience),
            "学历要求": str(education),
            "技能要求": str(skill_req),
            "岗位详情": str(job_detail),
            "招聘平台": str(detected_platform),
            "发布日期": str(publish_date),
            "跟进状态": "新线索",
            "抓取时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        if valid_url:
            feishu_data["岗位链接"] = {"link": valid_url, "text": "点击查看"}

        print(f">>> [3/3] 数据规整完毕，正在推送飞书...\n    公司: {company_name} | 岗位: {job_name} | 平台: {detected_platform}")
        success = feishu_api.create_feishu_record(feishu_data, TABLE_ID)
        print(f">>> 写入飞书结果: {'✅ 成功' if success else '❌ 失败'}")

        return success, feishu_data

    try:
        success, extracted_data = await asyncio.to_thread(process_llm_and_feishu, payload.raw_text)

        if success:
            print("====== 🎉 极速录入全流程完美收官！ ======\n")
            return {
                "status": "success",
                "message": "AI 解析并录入飞书成功！",
                "data": extracted_data
            }
        else:
            raise HTTPException(status_code=500, detail="大模型解析成功，但写入飞书时失败")

    except ValueError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        print(f"❌ 致命错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"极速录入失败: {str(e)}")

# ==================== 🌟 多图视觉录入 (Vision Import) API ====================

class JobImportImageRequest(BaseModel):
    images_base64: List[str]  # 接收前端传来的多张 base64 图片数组

@app.post("/api/jobs/import/image")
async def import_job_from_images(payload: JobImportImageRequest):
    import json
    print(f"\n====== 📡 收到【极速多图录入】请求，共 {len(payload.images_base64)} 张图片 ======")
    
    if not payload.images_base64:
        raise HTTPException(status_code=400, detail="没有接收到图片数据")

    system_prompt = """你是一个极度聪明的资深猎头助理。
用户会提供 1 张或多张包含招聘信息的截图。请仔细阅读这些截图中的文本，提取出核心字段，并严格按照以下 JSON 格式输出。

【提取规则】：
1. 必须输出纯净的 JSON 对象，绝不能包含 Markdown 代码块标记（如 ```json）或多余解释。
2. 如果多张图片的内容有重复，请自动去重拼凑；如果某些信息完全找不到，请填入 "未知" 或 "-"。
3. `招聘平台`字段：根据截图的 UI 界面（如小红书、Boss直聘、微信朋友圈等）判断来源并填入；若无法判断填 "截图解析"。
4. `岗位详情`字段：请尽可能完整地把图片里的岗位要求、职责等文字识别并拼接出来。

{
  "公司名称": "提取的公司名",
  "岗位名称": "提取的岗位名",
  "城市": "提取的城市",
  "薪资": "提取的薪资范围",
  "学历要求": "提取的学历要求",
  "经验要求": "提取的经验要求",
  "招聘平台": "判断的平台来源",
  "岗位详情": "提取到的完整 JD 描述"
}
"""

    def process_vision_and_feishu(base64_list):
        print(">>> [1/3] 正在建立与视觉大模型 (Vision Model) 的专属连接...")
        from openai import OpenAI
        from common import feishu_api
        from datetime import datetime
        # ⚠️ 注意：这里需要从 config 导入你的视觉模型名称
        from common.config import VISION_LLM_MODEL 
        
        temp_client = OpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL)
        
        # 构建多模态的 content 数组
        user_content = [{"type": "text", "text": "请解析以下截图中的招聘信息："}]
        for b64 in base64_list:
            # 确保前端传来的 base64 带有正确的前缀，如果没有我们需要自己补上
            img_url = b64 if b64.startswith("data:image") else f"data:image/jpeg;base64,{b64}"
            user_content.append({
                "type": "image_url",
                "image_url": {"url": img_url}
            })

        print(f">>> [2/3] 正在让大模型同时阅读 {len(base64_list)} 张图片... (Vision 模型耗时较长，请耐心等待)")
        response = temp_client.chat.completions.create(
            model=VISION_LLM_MODEL, # 🌟 使用独立的视觉模型
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ],
            temperature=0.1
        )
        
        content = (response.choices[0].message.content or "").strip()
        print(">>> ✅ 视觉大模型成功返回数据！开始解析 JSON...")
        
        prefix_json = "`" * 3 + "json"
        suffix = "`" * 3
        if content.startswith(prefix_json): content = content[7:]
        elif content.startswith(suffix): content = content[3:]
        if content.endswith(suffix): content = content[:-3]
        
        extracted_data = json.loads(content.strip())
        extracted_data["跟进状态"] = "新线索"
        extracted_data["抓取时间"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        print(">>> [3/3] 数据规整完毕，正在推送到飞书总表...")
        success = feishu_api.create_feishu_record(extracted_data, TABLE_ID)
        
        return success, extracted_data

    try:
        success, extracted_data = await asyncio.to_thread(process_vision_and_feishu, payload.images_base64)
        if success:
            print("====== 🎉 多图极速录入完美收官！ ======\n")
            return {"status": "success", "message": "图片解析并录入飞书成功！", "data": extracted_data}
        else:
            raise HTTPException(status_code=500, detail="大模型解析成功，但写入飞书时失败")
            
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="视觉大模型没有按格式返回 JSON，请重试")
    except Exception as e:
        print(f"❌ 多图解析致命错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"图片解析失败: {str(e)}")

@app.post("/api/strategy/test")
async def test_prompt_strategy(payload: PromptTestRequest):
    """根据前端传来的模式，动态套用不同的底层格式约束"""
    if not client:
        raise HTTPException(status_code=500, detail="AI 服务未配置")
        
    try:
        resume_text = await asyncio.to_thread(get_active_resume_from_feishu)
        if not resume_text:
            raise HTTPException(status_code=400, detail="未找到启用的简历底稿。")

        # 🌟 核心逻辑：根据测试模式，分配不同的“格式肉体”
        format_rules = ""
        
        if payload.test_mode == "evaluate":
            format_rules = """
            【输出格式要求（最高优先级）】
            必须、严格、仅输出一个纯净的 JSON 对象，不要包含任何 Markdown 代码块标记（如 ```json）。
            JSON 必须包含且仅包含以下字段：
            - "grade": 字符串，综合评级，必须是 A/B/C/D/F 之一（A=顶级匹配，B=良好，C=一般，D=较差，F=完全不匹配）
            - "scores": JSON 对象，包含以下 10 个键，每个值为 1-5 的整数：
                "role_match"（核心：角色与目标岗位的匹配程度），
                "skills_align"（核心：技能重合度），
                "seniority"（高权：职级资历匹配），
                "compensation"（高权：薪资期望契合度），
                "interview_prob"（高权：面试通过概率），
                "work_mode"（中权：工作模式契合，如远程/驻场），
                "company_stage"（中权：公司发展阶段契合度），
                "market_fit"（中权：赛道市场前景），
                "growth"（中权：个人成长空间），
                "timeline"（低权：招聘紧迫度与周期）
            - "extracted_skills": 字符串数组，提取 JD 明确要求的硬技能和工具
            - "dream_picture": 字符串。理想画像与能力信号总结
            - "risk_red_flags": 字符串。致命硬伤与后果推演
            - "deep_action_plan": 字符串。破局行动计划与关键信息索取
            """
        elif payload.test_mode == "rewrite":
            format_rules = """
            【必须严格遵守的排版与数据结构规范】（最高优先级）
            必须 100% 输出以下 JSON 结构，绝不能包含任何解释、前缀，不要写 ```json。
            {
              "personal_info": {"name": "", "contact": "", "city": "", "github_url": "", "portfolio_url": ""},
              "summary": "个人总结...",
              "skills": [{"category": "分类", "descriptions": ["技能"]}],
              "projects": [{"project_name": "项目名", "role": "角色", "time": "时间", "source_link": "链接", "tech_stack": "技术栈", "background": "**项目背景**：内容", "implementation": {"title": "**核心技术与落地**：", "points": ["动作"]}, "results": "**业务成果**：结果", "rewrite_rationale": "修改理由"}],
              "work_experience": [{"company_name": "公司", "title": "职位", "time": "时间", "experience_points": ["经历"], "rewrite_rationale": "理由"}],
              "education": [{"school": "学校", "degree": "学历", "time": "时间"}],
              "missing_data_requests": ["待办清单1"]
            }
            """
        else:
            format_rules = """
            【输出格式要求】
            请直接输出打招呼语的纯文本。绝不能包含任何代码块标记、思考过程、解释性语言或前缀（如“这是为您撰写的打招呼语：”）。
            """

        system_prompt = f"{payload.prompt_content}\n\n{format_rules}"
        user_prompt = f"【目标岗位 JD】\n{payload.jd_text}\n\n【原始简历】\n{resume_text}"

        response = await asyncio.to_thread(
            client.chat.completions.create,
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            timeout=2400.0
        )

        ai_result = (response.choices[0].message.content or "").strip()
        
        # 针对 JSON 模式的清理 (完整无截断的代码)
        if payload.test_mode in ["evaluate", "rewrite"]:
            if ai_result.startswith("```json"): 
                ai_result = ai_result[7:]
            elif ai_result.startswith("```"): 
                ai_result = ai_result[3:]
            if ai_result.endswith("```"): 
                ai_result = ai_result[:-3]
            
        return {"status": "success", "result": ai_result.strip()}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"沙盒运行异常: {str(e)}")

@app.post("/api/strategy/save")
def save_strategy_config(payload: SaveConfigRequest):
    """保存或更新配置到飞书多维表格"""
    try:
        token = get_tenant_access_token(APP_ID, APP_SECRET)
        table_id = FEISHU_TABLE_ID_RESUMES if payload.table_type == "resume" else FEISHU_TABLE_ID_PROMPTS
        
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
        
        try:
            if payload.record_id:
                # Update 存在的数据
                url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{table_id}/records/{payload.record_id}"
                resp = requests.put(url, headers=headers, json={"fields": payload.fields}, timeout=15, proxies={"http": None, "https": None}, verify=False)
            else:
                # Create 新增数据
                url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{table_id}/records"
                resp = requests.post(url, headers=headers, json={"fields": payload.fields}, timeout=15, proxies={"http": None, "https": None}, verify=False)
                
            data = resp.json()
            if resp.status_code == 200 and data.get("code") == 0:
                return {"status": "success", "record_id": data.get("data", {}).get("record", {}).get("record_id")}
                
            raise HTTPException(status_code=500, detail=f"保存失败: {data.get('msg')}")
        except requests.exceptions.SSLError as ssl_err:
            print(f"❌ SSL 错误 (save_strategy_config): {ssl_err}")
            raise HTTPException(status_code=502, detail=f"保存配置 SSL 错误，请检查网络代理设置: {str(ssl_err)}")
        except requests.exceptions.RequestException as req_err:
            print(f"❌ 网络请求错误 (save_strategy_config): {req_err}")
            raise HTTPException(status_code=502, detail=f"保存配置网络错误: {str(req_err)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理保存请求时出错: {e}")

# ============================================================================

# ==================== 🌟 新增：删除策略 (Delete) API ====================
class DeleteStrategyRequest(BaseModel):
    table_type: str # "resume" 或 "prompt"
    record_id: str

@app.post("/api/strategy/delete")
async def delete_strategy(payload: DeleteStrategyRequest):
    """从飞书中永久删除一条记录"""
    try:
        # 🌟 修复：直接从 config 中精准导入需要的变量，绝不报错
        from common.config import FEISHU_APP_TOKEN, FEISHU_TABLE_ID_PROMPTS, FEISHU_TABLE_ID_JOBS
        # 🌟 声明使用顶部已经导入好的全局 feishu_api
        global feishu_api 
        
        app_token = FEISHU_APP_TOKEN
        table_id = FEISHU_TABLE_ID_PROMPTS if payload.table_type == "prompt" else FEISHU_TABLE_ID_JOBS

        # 调用飞书删除记录 API
        success = feishu_api.delete_feishu_record(app_token, table_id, payload.record_id)
        
        if success:
            return {"status": "success"}
        else:
            raise HTTPException(status_code=500, detail="飞书 API 删除请求失败")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"删除失败: {str(e)}")


@app.post("/save_resume")
def save_resume(payload: SaveResumeRequest) -> Dict[str, Any]:
    app_id = _require_env("APP_ID", APP_ID)
    app_secret = _require_env("APP_SECRET", APP_SECRET)
    app_token = _require_env("APP_TOKEN", APP_TOKEN)
    table_id = _require_env("TABLE_ID", TABLE_ID)

    if not TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail=f"模板文件不存在: {TEMPLATE_PATH}")

    generated_image_paths: List[Path] = []
    try:
        tpl = DocxTemplate(str(TEMPLATE_PATH))
        tpl.render(payload.resume_data)
        tpl.save(str(TEMP_DOCX_PATH))

        convert(str(TEMP_DOCX_PATH), str(TEMP_PDF_PATH))
        if not TEMP_PDF_PATH.exists():
            raise HTTPException(status_code=500, detail="DOCX 转 PDF 失败，未生成 PDF 文件")

        pages = convert_from_path(str(TEMP_PDF_PATH), dpi=300, fmt="jpeg")
        if not pages:
            raise HTTPException(status_code=500, detail="PDF 转图片失败，未生成任何页面")

        for index, page in enumerate(pages, start=1):
            image_path = BASE_DIR / f"temp_resume_page_{index}.jpg"
            page.save(str(image_path), "JPEG", quality=95, optimize=True)
            generated_image_paths.append(image_path)

        tenant_access_token = get_tenant_access_token(app_id, app_secret)
        word_token = upload_file_to_feishu(tenant_access_token, app_token, TEMP_DOCX_PATH)
        pdf_token = upload_file_to_feishu(tenant_access_token, app_token, TEMP_PDF_PATH)
        image_tokens = [
            upload_file_to_feishu(tenant_access_token, app_token, image_path)
            for image_path in generated_image_paths
        ]

        update_bitable_record_attachments(
            tenant_access_token=tenant_access_token,
            app_token=app_token,
            table_id=table_id,
            record_id=payload.record_id,
            word_token=word_token,
            pdf_token=pdf_token,
            image_tokens=image_tokens,
        )

        return {
            "status": "success",
            "record_id": payload.record_id,
            "company_name": payload.company_name,
            "target_job": payload.target_job,
            "word_token": word_token,
            "pdf_token": pdf_token,
            "image_tokens": image_tokens,
        }
    except requests.RequestException as exc:
        raise HTTPException(status_code=502, detail=f"请求飞书接口失败: {exc}") from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"处理简历失败: {exc}") from exc
    finally:
        for path in [TEMP_DOCX_PATH, TEMP_PDF_PATH, *generated_image_paths]:
            if path.exists():
                try:
                    path.unlink()
                except OSError:
                    pass


@app.post("/ai_polish")
def ai_polish(payload: AIPolishRequest) -> Dict[str, Any]:
    if client is None:
        return {
            "polished_text": payload.selected_text,
            "error": "AI 服务未配置（缺少 api_key），请检查环境变量后重试。",
        }

    try:
        resp = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "你是一位顶级的资深 HR 和简历优化专家。请根据用户的要求，对提供的简历片段进行润色。必须严格遵守：只输出最终润色后的纯文本内容，绝对不要包含任何多余的解释、问候语，也绝对不要使用 markdown 代码块包裹（如 ``` 符号）。",
                },
                {
                    "role": "user",
                    "content": f"原始文本：{payload.selected_text}\n润色要求：{payload.instruction}",
                },
            ],
            timeout=30,
        )
        polished_text = (resp.choices[0].message.content or "").strip()
        if not polished_text:
            return {
                "polished_text": payload.selected_text,
                "error": "AI 返回为空，已保留原文本。",
            }
        return {"polished_text": polished_text}
    except Exception:
        return {
            "polished_text": payload.selected_text,
            "error": "AI 润色服务暂时不可用，请稍后重试。",
        }


@app.post("/export_resume")
def export_resume(payload: ExportRequest) -> Dict[str, Any]:
    import traceback
    from fastapi.responses import JSONResponse
    
    try:
        app_id = _require_env("APP_ID", APP_ID)
        app_secret = _require_env("APP_SECRET", APP_SECRET)
        app_token = _require_env("APP_TOKEN", APP_TOKEN)
        
        # 🌟 修改 3：统一写死使用总表 ID
        table_id = TABLE_ID

        export_type = (payload.export_type or "").lower().strip()
        if export_type not in {"word", "pdf", "image"}:
            raise HTTPException(status_code=400, detail="export_type 仅支持 word / pdf / image")

        # 🌟 动态模板选择：请求指定 > 启用模板 > 默认模板 > 兜底 template.docx
        template_path = EXPORT_TEMPLATE_PATH
        TEMPLATES_DIR.mkdir(exist_ok=True)
        if payload.template_name:
            _candidate = TEMPLATES_DIR / payload.template_name
            if _candidate.exists():
                template_path = _candidate
        else:
            _active_name = ""
            if ACTIVE_TEMPLATE_FILE.exists():
                _active_name = ACTIVE_TEMPLATE_FILE.read_text(encoding="utf-8").strip()
            if not _active_name and DEFAULT_TEMPLATE_MARKER.exists():
                _active_name = DEFAULT_TEMPLATE_MARKER.read_text(encoding="utf-8").strip()
            if _active_name:
                _candidate = TEMPLATES_DIR / _active_name
                if _candidate.exists():
                    template_path = _candidate
        if not template_path.exists():
            raise HTTPException(status_code=500, detail=f"模板不存在: {template_path}")

        record_id = extract_record_id(payload.job_id)
        image_paths: List[Path] = []

        tenant_access_token = get_tenant_access_token(app_id, app_secret)

        header = payload.resume_data.get("header", {})
        if not isinstance(header, dict):
            header = {}
        _raw_name = header.get("name", "候选人")
        if _raw_name is None or (isinstance(_raw_name, str) and not str(_raw_name).strip()):
            candidate_name = "候选人"
        else:
            candidate_name = str(_raw_name).strip() or "候选人"

        company_plain = "未知公司"
        job_plain = "未知岗位"
        platform_str = ""
        
        api_record = fetch_bitable_record_by_id(
            tenant_access_token, app_token, table_id, record_id
        )
        if api_record and isinstance(api_record.get("fields"), dict):
            fld = api_record["fields"]
            company_plain = feishu_field_to_plain_str(
                fld.get("公司名称"), "未知公司"
            ) or "未知公司"
            job_plain = feishu_field_to_plain_str(
                fld.get("岗位名称"), "未知岗位"
            ) or "未知岗位"
            platform_str = feishu_field_to_plain_str(fld.get("招聘平台"), "") or ""

        if not platform_str:
            platform_str = payload.platform or ""

        clean_name = sanitize_filename_component(candidate_name, "候选人")
        clean_company_name = sanitize_filename_component(company_plain, "未知公司")
        clean_job_title = sanitize_filename_component(job_plain, "未知岗位")

        if "猎聘" in platform_str:
            _combined = f"{clean_company_name}-{clean_job_title}"
            if len(_combined) <= 20:
                file_stem = sanitize_filename_component(_combined, clean_company_name)
            else:
                file_stem = sanitize_filename_component(clean_company_name, "未知公司")
        else:
            file_stem = f"{clean_name}-{clean_company_name}-{clean_job_title}简历"
        docx_filename = f"{file_stem}.docx"
        pdf_filename = f"{file_stem}.pdf"

        temp_docx = BASE_DIR / docx_filename
        temp_pdf = BASE_DIR / pdf_filename

        doc = DocxTemplate(str(template_path))
        
        context: Dict[str, Any] = {
            "name": header.get("name", ""),
            "contact": header.get("contact", ""),
            "intention": header.get("intention", ""),
            "sections": [],
        }

        for s in payload.resume_data.get("sections", []) or []:
            if not isinstance(s, dict):
                continue
            
            title = s.get("title", "")
            content = s.get("content", "")
            if not title.strip() and not content.strip():
                continue
            
            context["sections"].append(
                {
                    "title": title,
                    "content": parse_to_richtext(str(content), doc=doc),
                }
            )

        doc.render(context)
        doc.save(str(temp_docx))

        update_url = (
            f"https://open.feishu.cn/open-apis/bitable/v1/apps/{app_token}/tables/{table_id}/records/{record_id}"
        )
        update_headers = {"Authorization": f"Bearer {tenant_access_token}"}

        if export_type == "word":
            file_token = upload_file_to_feishu(
                tenant_access_token, app_token, temp_docx, file_name=docx_filename
            )
            fields: Dict[str, Any] = {"Word附件": [{"file_token": file_token}]}
            msg = "Word 已生成并写入「Word附件」"
        elif export_type == "pdf":
            convert(str(temp_docx), str(temp_pdf))
            if not temp_pdf.exists():
                raise HTTPException(status_code=500, detail="DOCX 转 PDF 失败")
            file_token = upload_file_to_feishu(
                tenant_access_token, app_token, temp_pdf, file_name=pdf_filename
            )
            fields = {"PDF备份": [{"file_token": file_token}]}
            msg = "PDF 已生成并写入「PDF备份」"
        else:
            convert(str(temp_docx), str(temp_pdf))
            if not temp_pdf.exists():
                raise HTTPException(status_code=500, detail="DOCX 转 PDF 失败（图片导出）")
            pages = convert_from_path(str(temp_pdf), dpi=300, fmt="jpeg")
            if not pages:
                raise HTTPException(status_code=500, detail="PDF 转图片失败")
            tokens: List[str] = []
            for index, page in enumerate(pages, start=1):
                img_filename = f"{file_stem}_p{index}.jpg"
                img_path = BASE_DIR / img_filename
                page.save(str(img_path), "JPEG", quality=95, optimize=True)
                image_paths.append(img_path)
                tokens.append(
                    upload_file_to_feishu(
                        tenant_access_token, app_token, img_path, file_name=img_filename
                    )
                )
            fields = {"图片保存": [{"file_token": t} for t in tokens]}
            msg = f"已生成 {len(tokens)} 张图片并写入「图片保存」"
        
        try:
            resp = requests.put(
                update_url,
                headers=update_headers,
                json={"fields": fields},
                timeout=15,
                proxies={"http": None, "https": None},
                verify=False
            )
            resp.raise_for_status()
            data = resp.json()
        except requests.exceptions.SSLError as ssl_err:
            print(f"❌ SSL 错误 (export_resume): {ssl_err}")
            raise HTTPException(
                status_code=502,
                detail=f"导出简历回写飞书 SSL 错误，请检查网络代理设置: {str(ssl_err)}",
            )
        except requests.exceptions.RequestException as req_err:
            print(f"❌ 网络请求错误 (export_resume): {req_err}")
            raise HTTPException(
                status_code=502,
                detail=f"导出简历回写飞书网络错误: {str(req_err)}",
            )
        
        if data.get("code") != 0:
            error_msg = data.get('msg', 'unknown error')
            raise HTTPException(
                status_code=502,
                detail=f"飞书回写失败: {error_msg}",
            )

        return {"status": "success", "msg": msg, "export_type": export_type, "record_id": record_id}

    except HTTPException as http_exc:
        raise
    except requests.RequestException as exc:
        raise HTTPException(status_code=502, detail=f"请求飞书接口失败: {exc}") from exc
    except Exception as exc:
        error_details = traceback.format_exc()
        return JSONResponse(
            status_code=500,
            content={
                "detail": f"后端内部错误: {str(exc)}",
                "error_type": type(exc).__name__,
                "traceback": error_details
            }
        )
    finally:
        for path in [temp_docx, temp_pdf, *image_paths]:
            if path.exists():
                try:
                    path.unlink()
                except OSError:
                    pass


@app.put("/api/update_job_status")
def update_job_status(payload: UpdateJobStatusRequest) -> Dict[str, Any]:
    try:
        pure_record_id = extract_record_id(payload.job_id)
        
        # 🌟 修改 4：统一更新状态也走总表
        success = feishu_api.update_feishu_record(
            record_id=pure_record_id,
            fields_to_update={"跟进状态": payload.status},
            table_id=TABLE_ID
        )
        
        if success:
            return {
                "status": "success",
                "message": f"成功更新跟进状态为: {payload.status}",
                "job_id": pure_record_id,
                "updated_status": payload.status
            }
        else:
            raise HTTPException(
                status_code=500,
                detail="更新飞书跟进状态失败"
            )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"更新跟进状态时发生错误: {str(e)}"
        )


@app.put("/api/update_review_comments")
def update_review_comments(payload: UpdateReviewCommentsRequest) -> Dict[str, Any]:
    try:
        pure_record_id = extract_record_id(payload.job_id)
        
        success = feishu_api.update_feishu_record(
            record_id=pure_record_id,
            fields_to_update={"我的复核": payload.comments}
        )
        
        if success:
            return {
                "status": "success",
                "message": "成功更新复核意见",
                "job_id": pure_record_id,
                "updated_comments": payload.comments
            }
        else:
            raise HTTPException(
                status_code=500,
                detail="更新飞书复核意见失败"
            )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"更新复核意见时发生错误: {str(e)}"
        )


@app.post("/api/jobs/batch-delete")
def batch_delete_jobs(payload: BatchDeleteRequest):
    try:
        if not payload.job_ids:
            return {"status": "success", "message": "没有需要删除的记录"}

        pure_record_ids = [extract_record_id(jid) for jid in payload.job_ids]

        success = feishu_api.batch_delete_feishu_records(pure_record_ids, TABLE_ID)

        if success:
            return {"status": "success", "message": f"成功删除 {len(pure_record_ids)} 个岗位"}
        else:
            raise HTTPException(status_code=500, detail="飞书批量删除失败")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"批量删除时发生错误: {str(e)}")


@app.put("/api/update_greeting")
def update_greeting(payload: UpdateGreetingRequest) -> Dict[str, Any]:
    try:
        pure_record_id = extract_record_id(payload.job_id)

        success = feishu_api.update_feishu_record(
            record_id=pure_record_id,
            fields_to_update={"打招呼语": payload.greeting},
            table_id=TABLE_ID
        )

        if success:
            return {
                "status": "success",
                "message": "成功更新打招呼语",
                "job_id": pure_record_id,
                "updated_greeting": payload.greeting
            }
        else:
            raise HTTPException(
                status_code=500,
                detail="更新飞书打招呼语失败"
            )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"更新打招呼语时发生错误: {str(e)}"
        )


@app.post("/api/save_manual_resume")
def save_manual_resume(payload: SaveManualResumeRequest) -> Dict[str, Any]:
    try:
        pure_record_id = extract_record_id(payload.job_id)
        
        # 🌟 修改 5：统一人工保存走总表
        success = feishu_api.update_qa_fields(
            record_id=pure_record_id,
            manual_refined_resume=payload.resume_text,
            table_id=TABLE_ID
        )
        
        if success:
            return {
                "status": "success",
                "message": "简历已保存到飞书",
                "job_id": payload.job_id
            }
        else:
            raise HTTPException(
                status_code=500,
                detail="保存简历到飞书失败"
            )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"保存简历失败: {str(e)}"
        )


# ─── 模板仓库 API ─────────────────────────────────────────────────────────────

def scan_docx_variables(file_path: Path) -> Dict[str, List[str]]:
    """双通道扫描 .docx 模板变量和逻辑标签，彻底消除 Word Run 碎片割裂问题。

    Pass 1 — python-docx .text（无分隔符拼接）：
        paragraph.text 自动合并同段落内所有 run；"".join() 无间隔拼接各段落，
        确保 {{ 在段尾、变量名在段首这类跨段碎片也能被正则命中。
    Pass 2 — 原始 XML 兜底（doc.element.xml）：
        覆盖文本框、自定义 XML 等 pass1 可能遗漏的区域。
    """
    from docx import Document as _DocxDoc

    # 宽松正则：{{ }} 内允许任意空白
    var_re = re.compile(r'\{\{-?\s*([\w]+(?:\.[\w]+)*)\s*-?\}\}')
    # 捕获 {% for item in sections %} 中的集合变量
    for_var_re = re.compile(r'\{%-?\s*for\s+\w+\s+in\s+([\w]+)')
    # 捕获指令关键字（for / if / endif 等）
    tag_kw_re = re.compile(r'\{%-?\s*(\w+)')

    variables: set = set()
    tag_keywords: set = set()

    def _scan(text: str) -> None:
        for m in var_re.finditer(text):
            variables.add(m.group(1).split('.')[0])
        for m in for_var_re.finditer(text):
            variables.add(m.group(1).split('.')[0])
        for m in tag_kw_re.finditer(text):
            tag_keywords.add(m.group(1))

    try:
        doc = _DocxDoc(str(file_path))

        # ── Pass 1: 无间隔拼接所有段落文本，消除跨段落碎片 ──────────────────
        parts: List[str] = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        parts.append(para.text)
        for section in doc.sections:
            for hdr in (section.header, section.footer):
                if hdr:
                    for para in hdr.paragraphs:
                        parts.append(para.text)
        _scan("".join(parts))  # ← 无分隔符，{{ 和 }} 不会因换行被撕裂

        # ── Pass 2: 原始 XML 宽松扫描，兜底捕获文本框等遗漏区域 ─────────────
        raw_xml = doc.element.xml
        xml_text = re.sub(r'<[^>]+>', ' ', raw_xml)   # 标签 → 单空格
        xml_text = re.sub(r'\s+', ' ', xml_text)       # 合并多余空白
        _scan(xml_text)

    except Exception as e:
        print(f"⚠️ scan_docx_variables 失败 [{file_path.name}]: {e}")

    return {"variables": sorted(variables), "tags": sorted(tag_keywords)}


def generate_template_preview(filename: str) -> bool:
    """将 .docx 转换为 PDF，再截取第一页为 JPG 缩略图。失败时静默返回 False。"""
    try:
        docx_path = TEMPLATES_DIR / filename
        if not docx_path.exists():
            return False
        stem = Path(filename).stem
        pdf_path = TEMPLATES_DIR / f"{stem}_preview_tmp.pdf"
        preview_path = TEMPLATES_DIR / f"preview_{stem}.jpg"

        if preview_path.exists():  # 已有最新预览图，跳过重新生成
            return True

        convert(str(docx_path), str(pdf_path))

        images = convert_from_path(
            str(pdf_path),
            first_page=1,
            last_page=1,
            dpi=130,
        )
        if images:
            images[0].save(str(preview_path), "JPEG", quality=85)

        if pdf_path.exists():
            pdf_path.unlink()

        print(f"✅ 模板预览图已生成: {preview_path.name}")
        return True
    except Exception as e:
        print(f"⚠️ generate_template_preview 失败 [{filename}]: {e}")
        return False


@app.get("/api/templates")
def list_templates() -> Dict[str, Any]:
    TEMPLATES_DIR.mkdir(exist_ok=True)
    default_name = ""
    if DEFAULT_TEMPLATE_MARKER.exists():
        default_name = DEFAULT_TEMPLATE_MARKER.read_text(encoding="utf-8").strip()
    files = sorted(TEMPLATES_DIR.glob("*.docx"))
    result = []
    for f in files:
        scan = scan_docx_variables(f)
        result.append({
            "name": f.name,
            "size": f.stat().st_size,
            "is_default": f.name == default_name,
            "variables": scan["variables"],
            "tags": scan["tags"],
        })
    if result and not default_name:
        result[0]["is_default"] = True
        default_name = result[0]["name"]
    return {"status": "success", "templates": result, "default": default_name}


@app.get("/api/templates/active")
def get_active_template() -> Dict[str, Any]:
    """返回当前启用模板文件名。优先级: active_template.txt > .default > 第一个可用模板。"""
    TEMPLATES_DIR.mkdir(exist_ok=True)
    active_name = ""
    if ACTIVE_TEMPLATE_FILE.exists():
        active_name = ACTIVE_TEMPLATE_FILE.read_text(encoding="utf-8").strip()
    if not active_name and DEFAULT_TEMPLATE_MARKER.exists():
        active_name = DEFAULT_TEMPLATE_MARKER.read_text(encoding="utf-8").strip()
    if not active_name:
        files = sorted(TEMPLATES_DIR.glob("*.docx"))
        if files:
            active_name = files[0].name
    return {"status": "success", "active": active_name}


@app.post("/api/templates/active/{template_name}")
def set_active_template_route(template_name: str) -> Dict[str, Any]:
    """将指定模板设为全局启用状态（排他其他）。"""
    TEMPLATES_DIR.mkdir(exist_ok=True)
    target = TEMPLATES_DIR / template_name
    if not target.exists():
        raise HTTPException(status_code=404, detail=f"模板文件不存在: {template_name}")
    ACTIVE_TEMPLATE_FILE.write_text(template_name, encoding="utf-8")
    return {"status": "success", "active": template_name}


@app.get("/api/templates/download/{template_name}")
def download_template(template_name: str) -> FileResponse:
    TEMPLATES_DIR.mkdir(exist_ok=True)
    target = TEMPLATES_DIR / template_name
    if not target.exists():
        raise HTTPException(status_code=404, detail=f"模板文件不存在: {template_name}")
    return FileResponse(
        path=str(target),
        filename=template_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/api/templates/preview/{template_name}")
def get_template_preview(template_name: str):
    """返回模板第一页预览图（JPG），不存在时返回 404。"""
    stem = Path(template_name).stem
    preview_path = TEMPLATES_DIR / f"preview_{stem}.jpg"
    if not preview_path.exists():
        raise HTTPException(status_code=404, detail="预览图尚未生成，请稍后再试")
    return FileResponse(path=str(preview_path), media_type="image/jpeg")


@app.post("/api/templates")
async def upload_template(background_tasks: BackgroundTasks, file: UploadFile = File(...)) -> Dict[str, Any]:
    TEMPLATES_DIR.mkdir(exist_ok=True)
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .docx 格式")
    safe_name = re.sub(r'[^\w\-_. \u4e00-\u9fff]', '_', file.filename)
    target_path = TEMPLATES_DIR / safe_name
    old_preview = TEMPLATES_DIR / f"preview_{Path(safe_name).stem}.jpg"
    if old_preview.exists():
        old_preview.unlink()
    content = await file.read()
    with open(target_path, "wb") as buf:
        buf.write(content)
    scan = scan_docx_variables(target_path)
    background_tasks.add_task(generate_template_preview, safe_name)
    return {"status": "success", "name": safe_name, "variables": scan["variables"], "tags": scan["tags"]}


@app.put("/api/templates/{template_name}/default")
def set_default_template(template_name: str) -> Dict[str, Any]:
    TEMPLATES_DIR.mkdir(exist_ok=True)
    target = TEMPLATES_DIR / template_name
    if not target.exists():
        raise HTTPException(status_code=404, detail=f"模板文件不存在: {template_name}")
    DEFAULT_TEMPLATE_MARKER.write_text(template_name, encoding="utf-8")
    return {"status": "success", "default": template_name}


@app.delete("/api/templates/{template_name}")
def delete_template(template_name: str) -> Dict[str, Any]:
    TEMPLATES_DIR.mkdir(exist_ok=True)
    target = TEMPLATES_DIR / template_name
    if not target.exists():
        raise HTTPException(status_code=404, detail=f"模板文件不存在: {template_name}")
    target.unlink()
    if DEFAULT_TEMPLATE_MARKER.exists():
        if DEFAULT_TEMPLATE_MARKER.read_text(encoding="utf-8").strip() == template_name:
            DEFAULT_TEMPLATE_MARKER.unlink()
    return {"status": "success"}


@app.put("/api/templates/{template_name}")
async def replace_template(template_name: str, background_tasks: BackgroundTasks, file: UploadFile = File(...)) -> Dict[str, Any]:
    """覆盖上传指定模板（保持原文件名，替换内容）"""
    TEMPLATES_DIR.mkdir(exist_ok=True)
    target = TEMPLATES_DIR / template_name
    if not target.exists():
        raise HTTPException(status_code=404, detail=f"模板文件不存在: {template_name}")
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .docx 格式")
    old_preview = TEMPLATES_DIR / f"preview_{Path(template_name).stem}.jpg"
    if old_preview.exists():
        old_preview.unlink()
    content = await file.read()
    with open(target, "wb") as buf:
        buf.write(content)
    scan = scan_docx_variables(target)
    background_tasks.add_task(generate_template_preview, template_name)
    return {"status": "success", "name": template_name, "variables": scan["variables"], "tags": scan["tags"]}


def compress_jd_text(jd_text: str, max_length: int = 800) -> str:
    if not jd_text or len(jd_text) <= max_length:
        return jd_text
    
    compressed = jd_text[:max_length]
    return compressed + "\n\n...(后续内容已省略)"


def flatten_resume_json(resume_json_str: str) -> str:
    if not resume_json_str:
        return ""
    
    try:
        import json
        resume_data = json.loads(resume_json_str)
        
        flattened_parts = []
        
        if isinstance(resume_data, dict) and "header" in resume_data:
            header = resume_data["header"]
            if isinstance(header, dict):
                name = header.get("name", "")
                intention = header.get("intention", "")
                if name or intention:
                    flattened_parts.append(f"# {name} - {intention}")
        
        if isinstance(resume_data, dict) and "sections" in resume_data:
            sections = resume_data["sections"]
            if isinstance(sections, list):
                for section in sections:
                    if isinstance(section, dict):
                        title = section.get("title", "")
                        content = section.get("content", "")
                        
                        if title and content:
                            flattened_parts.append(f"\n## {title}\n{content}")
        
        if not flattened_parts:
            if isinstance(resume_data, dict):
                for key, value in resume_data.items():
                    if isinstance(value, str) and len(value) > 20:
                        if key not in ["id", "record_id", "rewrite_rationale", "rationale"]:
                            flattened_parts.append(f"\n{value}")
        
        result = "\n".join(flattened_parts).strip()
        
        if len(result) > 3000:
            result = result[:3000] + "\n\n...(后续内容已省略)"
        
        return result
        
    except json.JSONDecodeError:
        if len(resume_json_str) > 3000:
            return resume_json_str[:3000] + "\n\n...(后续内容已省略)"
        return resume_json_str
    except Exception as e:
        return resume_json_str[:500] if resume_json_str else ""


@app.post("/copilot_chat")
def copilot_chat(payload: CopilotChatRequest) -> Dict[str, Any]:
    if client is None:
        raise HTTPException(
            status_code=500,
            detail="AI 服务未配置（缺少 LLM_API_KEY），请检查环境变量后重试。"
        )
    
    try:
        jd_text_raw = payload.context.get("jd_text", "")
        evaluation_report = payload.context.get("evaluation_report", "")
        ai_resume_json = payload.context.get("ai_resume_json", "")
        human_refined_resume = payload.context.get("human_refined_resume", "")
        
        jd_text = compress_jd_text(jd_text_raw, max_length=800)
        
        current_resume_raw = human_refined_resume if human_refined_resume else ai_resume_json
        current_resume = flatten_resume_json(current_resume_raw)
        
        system_prompt = f"""你是一位拥有 15 年经验的顶级求职 Copilot，专注于帮助候选人深度理解岗位需求、优化简历策略、准备面试答辩。

【当前上下文】（已压缩优化）：

**目标岗位 JD（核心摘要）**：
{jd_text if jd_text else "暂无岗位详情"}

**AI 评估报告**：
{evaluation_report if evaluation_report else "暂无评估报告"}

**当前简历数据（扁平化）**：
{current_resume if current_resume else "暂无简历数据"}

【沟通准则】：
- 回答必须专业、精准、有深度，避免套话和泛泛而谈
- 引用具体的 JD 片段、简历内容、评估结论来支撑你的建议
- 保持简洁，每次回复控制在 200-300 字，除非用户明确要求详细解答
"""
        
        messages = [{"role": "system", "content": system_prompt}]
        
        for msg in payload.history:
            if msg.get("role") and msg.get("content"):
                messages.append({
                    "role": msg["role"],
                    "content": msg["content"]
                })
        
        messages.append({
            "role": "user",
            "content": payload.user_question
        })
        
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=messages,
            temperature=0.7,
            timeout=60.0, 
        )
        
        ai_reply = (response.choices[0].message.content or "").strip()
        
        if not ai_reply:
            raise HTTPException(
                status_code=500,
                detail="AI 返回为空，请稍后重试"
            )
        
        return {
            "status": "success",
            "reply": ai_reply
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Copilot Chat 服务暂时不可用: {str(e)}"
        )


@app.post("/api/qa_evaluate")
def qa_evaluate(payload: QAEvaluateRequest) -> Dict[str, Any]:
    pure_record_id = extract_record_id(payload.job_id)
    
    try:
        qa_report = qa_evaluate_resume(
            job_description=payload.job_description,
            rewritten_resume_text=payload.resume_text
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"LLM 调用失败: {str(e)}"
        )
    
    if not qa_report:
        raise HTTPException(
            status_code=500,
            detail="LLM 返回了空的质检报告"
        )
    
    try:
        # 🌟 修改 6：统一 QA 回写走总表
        qa_success = feishu_api.update_qa_fields(
            record_id=pure_record_id,
            qa_report_json=qa_report,
            table_id=TABLE_ID
        )
        
        return {
            "status": "success",
            "message": "QA 评估完成",
            "job_id": payload.job_id,
            "qa_report_saved": qa_success,
            "qa_report": qa_report
        }
    except Exception as e:
        return {
            "status": "success",
            "message": "QA 评估完成，但保存到飞书失败",
            "job_id": payload.job_id,
            "qa_report_saved": False,
            "qa_report": qa_report
        }


def get_job_record_from_feishu(record_id: str, table_id: str) -> Optional[Dict[str, Any]]:
    token = feishu_api.get_tenant_access_token()
    if not token:
        return None
    
    url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{table_id}/records/{record_id}"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    
    try:
        response = requests.get(url, headers=headers, timeout=15, proxies={"http": None, "https": None}, verify=False)
        data = response.json()
        if data.get("code") == 0:
            return data.get("data", {}).get("record", {})
        else:
            return None
    except requests.exceptions.SSLError as ssl_err:
        print(f"❌ SSL 错误 (get_job_record_from_feishu): {ssl_err}")
        return None
    except requests.exceptions.RequestException as req_err:
        print(f"❌ 网络请求错误 (get_job_record_from_feishu): {req_err}")
        return None
    except Exception as e:
        print(f"❌ 获取岗位记录失败: {e}")
        return None


async def send_sse_msg(queue: asyncio.Queue, msg_type: str, message: str, **kwargs):
    payload = {"type": msg_type, "message": message}
    payload.update(kwargs)
    await queue.put(f"data: {json.dumps(payload, ensure_ascii=False)}\n\n")


async def run_batch_ai_task(
    task_id: str,
    task_type: str,
    job_ids: List[str],
    queue: asyncio.Queue,
    scheduled_at: Optional[str] = None
):
    GLOBAL_TASK_STATE["is_processing"] = True  # 🔒 任务开始，上锁
    try:
        task_status[task_id]["status"] = "running"
        task_status[task_id]["started_at"] = datetime.now().isoformat()
        
        await send_sse_msg(queue, "start", f"批量任务已启动，共 {len(job_ids)} 个岗位", total=len(job_ids))
        
        # 🌟 核心修改：现在直接从飞书拉取【云端简历】！再也不用看本地的 txt 脸色了
        await send_sse_msg(queue, "info", "☁️ 正在从飞书配置中心读取【启用】状态的简历...")
        resume_text = await asyncio.to_thread(get_active_resume_from_feishu)
        
        if not resume_text:
            raise Exception("未在飞书【我的简历库】中找到处于[启用]状态的简历！请前往控制台添加。")
        
        await send_sse_msg(queue, "info", f"✓ 简历文本已加载，共 {len(resume_text)} 字符")
        
        await send_sse_msg(queue, "info", "⚙️ 正在加载求职偏好与底线配置...")
        try:
            if hasattr(feishu_api, "get_my_preferences"):
                preferences_text = await asyncio.to_thread(feishu_api.get_my_preferences)
            else:
                print("⚠️ feishu_api 中未找到 get_my_preferences，已降级为空偏好")
                preferences_text = ""
        except Exception as pref_err:
            print(f"⚠️ 读取求职偏好失败，已降级为空偏好继续执行: {pref_err}")
            preferences_text = ""
        
        for index, job_id in enumerate(job_ids, start=1):
            print(f"\n🔍 [正在处理] 第 {index} 个岗位: {job_id}")
            _sent_terminal = False
            job_updates = {}
            try:
                await send_sse_msg(queue, "progress", f"正在处理 {index}/{len(job_ids)}: {job_id}...", job_id=job_id, current=index, total=len(job_ids))
                
                record_id = extract_record_id(job_id)
                
                # 🌟 修改 7：后台任务强制使用总表 ID
                table_id = TABLE_ID
                platform = "未知"
                if "猎聘" in job_id:
                    platform = "猎聘"
                elif "51job" in job_id or "前程无忧" in job_id:
                    platform = "51job"
                elif "智联招聘" in job_id:
                    platform = "智联招聘"
                else:
                    platform = "BOSS直聘"
                
                await send_sse_msg(queue, "info", f"  📡 正在从飞书拉取【{platform}】岗位详情...")
                
                job_record = await asyncio.to_thread(
                    get_job_record_from_feishu,
                    record_id,
                    table_id
                )
                
                if not job_record:
                    raise Exception(f"无法从飞书获取岗位记录 {record_id}")
                
                fields = job_record.get("fields", {})
                job_name = feishu_api.extract_feishu_text(fields.get("岗位名称", ""))
                company_name = feishu_api.extract_feishu_text(fields.get("公司名称", ""))
                jd_text = feishu_api.extract_feishu_text(fields.get("岗位详情", ""))
                salary = feishu_api.extract_feishu_text(fields.get("薪资", ""))
                city = feishu_api.extract_feishu_text(fields.get("城市", ""))
                experience = feishu_api.extract_feishu_text(fields.get("经验要求", ""))
                education = feishu_api.extract_feishu_text(fields.get("学历要求", ""))
                
                print(f"📡 [飞书数据] 已成功拉取到 {company_name} - {job_name} 的岗位记录")
                await send_sse_msg(queue, "info", f"  ✓ 已获取【{company_name} - {job_name}】的岗位详情")
                
                if not jd_text or len(jd_text.strip()) < 50:
                    print(f"⚠️ [跳过] 岗位 {job_id} 详情长度不足({len(jd_text) if jd_text else 0})，无法处理。")
                    raise Exception("岗位详情为空或过短，无法进行 AI 处理")
                
                if task_type == "evaluate":
                    await send_sse_msg(queue, "info", f"  🔍 正在背调【{company_name}】公司情报...")
                    company_intel = await research_company_serper(company_name)
                    await send_sse_msg(queue, "info", "  🧠 正在呼叫 LLM 进行深度评估...")
                    
                    job_data = {
                        "record_id": record_id,
                        "table_id": table_id,
                        "platform": platform,
                        "company": company_name,
                        "job_title": job_name,
                        "jd_text": jd_text,
                        "salary": salary,
                        "city": city,
                        "experience": experience,
                        "education": education,
                    }
                    
                    import time
                    start_time = time.time()
                    
                    try:
                        result = await asyncio.to_thread(
                            evaluate_single_job,
                            job_data,
                            resume_text,
                            company_intel,
                            preferences_text
                        )
                    except ValueError:
                        await send_sse_msg(queue, "error", "  ⚠️ AI 返回格式异常，已触发自动纠偏或跳过", job_id=job_id)
                        task_status[task_id]["completed"] = index
                        _sent_terminal = True
                        continue

                    elapsed_time = time.time() - start_time

                    if result["success"]:
                        ai_score = result.get("ai_score", 0)
                        target_status = result.get("status", "待人工评估")
                        usage = result.get("usage", {})
                        pt = usage.get("prompt_tokens", 0)
                        ct = usage.get("completion_tokens", 0)
                        tt = usage.get("total_tokens", 0)

                        await send_sse_msg(queue, "info", f"📈 Token 消耗 → 提示: {pt} / 补全: {ct}", job_id=job_id)
                        await send_sse_msg(queue, "info", f"  📊 诊断报告已出！AI 综合得分: {ai_score}分 (耗时 {elapsed_time:.1f}s)", usage={"prompt": pt, "completion": ct, "total": tt})
                        await send_sse_msg(queue, "info", "  ☁️ 正在将评估结果回写到飞书...")
                        
                        update_success = await asyncio.to_thread(
                            feishu_api.update_feishu_record,
                            record_id,
                            result["update_data"],
                            table_id
                        )
                        
                        if update_success:
                            job_updates = {"followStatus": target_status, "aiScore": ai_score, "grade": result.get("grade", "")}
                            await send_sse_msg(queue, "info", f"  ✅ 评估完成，状态已更新为: {target_status}", usage={"prompt": pt, "completion": ct, "total": tt})
                            rationales_text = result.get("rationales_text", "")
                            if rationales_text:
                                for dim_block in rationales_text.split("\n\n")[:10]:
                                    await send_sse_msg(queue, "info", f"[诊断] {dim_block}")
                        else:
                            await send_sse_msg(queue, "warning", "  ⚠️ 评估完成，但回写飞书失败")
                    else:
                        raise Exception(result.get("error", "评估失败"))

                elif task_type == "rewrite":
                    await send_sse_msg(queue, "info", "  🧠 正在构思高情商打招呼语...")
                    await send_sse_msg(queue, "info", "  📝 正在根据 JD 生成 Markdown 定制简历...")

                    import time
                    start_time = time.time()

                    # 🌟 核心修复：把飞书里所有的深度评估精华结论统统喂给改写引擎
                    diagnosis_dict = {
                        "综合评级": feishu_api.extract_feishu_text(fields.get("综合评级 (A-F)", "")),
                        "AI评估详情": feishu_api.extract_feishu_text(fields.get("AI评估详情", "")),
                        "理想画像与能力信号": feishu_api.extract_feishu_text(fields.get("理想画像与能力信号", "")),
                        "核心能力词典": feishu_api.extract_feishu_text(fields.get("核心能力词典", "")),
                        "高杠杆匹配点": feishu_api.extract_feishu_text(fields.get("高杠杆匹配点", "")),
                        "致命硬伤与毒点": feishu_api.extract_feishu_text(fields.get("致命硬伤与毒点", "")),
                        "破局行动计划": feishu_api.extract_feishu_text(fields.get("破局行动计划", "")),
                    }

                    try:
                        md_resume, rewrite_usage = await asyncio.to_thread(
                            rewrite_resume_for_job,
                            resume_text,
                            jd_text,
                            diagnosis_dict,
                            job_name
                        )
                    except ValueError:
                        await send_sse_msg(queue, "error", "  ⚠️ AI 返回格式异常，已触发自动纠偏或跳过", job_id=job_id)
                        task_status[task_id]["completed"] = index
                        _sent_terminal = True
                        continue

                    greeting, greeting_usage = await asyncio.to_thread(
                        generate_greeting,
                        resume_text,
                        jd_text,
                        job_name
                    )

                    elapsed_time = time.time() - start_time
                    
                    # 🌟 修复 1：把“简历改写”和“打招呼语”两次调用大模型的 Token 加起来，计算才准确
                    rw_pt = (rewrite_usage.get("prompt_tokens", 0) if rewrite_usage else 0) + (greeting_usage.get("prompt_tokens", 0) if greeting_usage else 0)
                    rw_ct = (rewrite_usage.get("completion_tokens", 0) if rewrite_usage else 0) + (greeting_usage.get("completion_tokens", 0) if greeting_usage else 0)
                    rw_tt = (rewrite_usage.get("total_tokens", 0) if rewrite_usage else 0) + (greeting_usage.get("total_tokens", 0) if greeting_usage else 0)
                    
                    await send_sse_msg(queue, "info", f"📈 Token 消耗 → 提示: {rw_pt} / 补全: {rw_ct} / 总计: {rw_tt}", job_id=job_id)

                    if md_resume and greeting:
                        sections = parse_resume_markdown(md_resume)
                        # 注意：这里我们刚刚修复过 Token 的代码，如果你刚才已经替换了 Token 逻辑，只需确保 fields_to_update 和 job_updates 如下修改即可
                        await send_sse_msg(
                            queue, 
                            "info", 
                            f"  ✓ 改写完成，解析出 {len(sections)} 个模块 (耗时 {elapsed_time:.1f}s)", 
                            usage={"prompt": rw_pt, "completion": rw_ct, "total": rw_tt}
                        )

                        fields_to_update = {
                            "打招呼语": greeting,
                            "AI改写JSON": md_resume,
                            "跟进状态": "简历人工复核"  # 🌟 修复 1：明确告知飞书更新状态
                        }

                        await send_sse_msg(queue, "info", "  ☁️ 正在将打招呼语与 AI 改写 Markdown 写回飞书...")

                        update_success = await asyncio.to_thread(
                            feishu_api.update_feishu_record,
                            record_id,
                            fields_to_update,
                            table_id
                        )

                        if update_success:
                            job_updates = {"followStatus": "简历人工复核"}  # 🌟 修复 2：前端 UI 状态同步纠正
                            await send_sse_msg(queue, "info", "  ✅ 简历改写完成，状态已更新为: 简历人工复核")
                        else:
                            await send_sse_msg(queue, "warning", "  ⚠️ 改写完成，但回写飞书失败")
                    else:
                        raise Exception("AI 生成打招呼语或简历失败")

                elif task_type == "deep_rewrite":
                    await send_sse_msg(queue, "info", "🧠 正在召集多Agent专家组进行深度研讨...")

                    import time
                    start_time = time.time()

                    # 🌟 1. 提取之前评估写在飞书里的诊断数据
                    diagnosis_dict = {
                        "综合评级": feishu_api.extract_feishu_text(fields.get("综合评级 (A-F)", "")),
                        "理想画像与能力信号": feishu_api.extract_feishu_text(fields.get("理想画像与能力信号", "")),
                        "核心能力词典": feishu_api.extract_feishu_text(fields.get("核心能力词典", "")),
                        "高杠杆匹配点": feishu_api.extract_feishu_text(fields.get("高杠杆匹配点", "")),
                        "致命硬伤与毒点": feishu_api.extract_feishu_text(fields.get("致命硬伤与毒点", "")),
                    }
                    import json
                    diagnosis_report_str = json.dumps(diagnosis_dict, ensure_ascii=False, indent=2)

                    # 🌟 2. 注入符合最新 ResumeState 的 initial_state
                    initial_state = {
                        "original_full_text": resume_text,
                        "jd": jd_text,
                        "diagnosis_report": diagnosis_report_str,
                        "parsed_blocks": [],
                        "working_rewritten_blocks": [],
                        "critic_feedback": "",
                        "current_score": 0,
                        "revision_count": 0,
                        "token_usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0},
                        "logs": [],
                        "final_markdown": ""
                    }

                    md_resume = ""
                    final_token_usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
                    printed_log_count = 0  # 🌟 新增：追踪已打印的日志数量，防止漏打
                    
                    try:
                        async for output in multi_agent_app.astream(initial_state):
                            for node_name, state_update in output.items():
                                if "token_usage" in state_update:
                                    final_token_usage = state_update["token_usage"]

                                # 🌟 核心修复：遍历所有新增加的 logs，确保终端和前端都能完整看到 Agent 的每一步思考
                                current_logs = state_update.get("logs", [])
                                while printed_log_count < len(current_logs):
                                    log_msg = current_logs[printed_log_count]
                                    await send_sse_msg(queue, "info", f"💬 {log_msg}")
                                    print(f"💬 [多Agent监控] {log_msg}")
                                    printed_log_count += 1

                                # 🌟 匹配新的字段名 final_markdown
                                if "final_markdown" in state_update:
                                    md_resume = state_update["final_markdown"]
                    except Exception as agent_err:
                        print(f"❌ [多Agent执行崩溃]: {str(agent_err)}")
                        raise Exception(f"多Agent工作流执行失败: {agent_err}")

                    if not md_resume:
                        raise Exception("多Agent工作流未能生成最终简历")

                    greeting, _ = await asyncio.to_thread(
                        generate_greeting,
                        resume_text,
                        jd_text,
                        job_name
                    )

                    elapsed_time = time.time() - start_time
                    pt = final_token_usage.get("prompt_tokens", 0)
                    ct = final_token_usage.get("completion_tokens", 0)
                    tt = final_token_usage.get("total_tokens", 0)

                    print(f"📈 [多Agent Token消耗] 提示: {pt} / 补全: {ct} / 总计: {tt}")
                    await send_sse_msg(queue, "info", f"📈 Token 消耗 → 提示: {pt} / 补全: {ct} / 总计: {tt}", job_id=job_id)
                    await send_sse_msg(queue, "info", f"  ✓ 多Agent改写完成 (耗时 {elapsed_time:.1f}s)", usage={"prompt": pt, "completion": ct, "total": tt})

                    fields_to_update = {
                        "AI改写JSON": md_resume,
                        "打招呼语": greeting,
                        "跟进状态": "简历人工复核"  # 🌟 修复：将其从"待人工复核"统一改为"简历人工复核"
                    }

                    await send_sse_msg(queue, "info", "  ☁️ 正在将多Agent改写结果回写到飞书...")

                    update_success = await asyncio.to_thread(
                        feishu_api.update_feishu_record,
                        record_id,
                        fields_to_update,
                        table_id
                    )

                    if update_success:
                        job_updates = {"followStatus": "简历人工复核"}
                        await send_sse_msg(queue, "info", "  ✅ 多Agent改写完成，状态已更新为: 简历人工复核")
                    else:
                        await send_sse_msg(queue, "warning", "  ⚠️ 改写完成，但回写飞书失败")

                elif task_type == "deep_evaluate":
                    await send_sse_msg(queue, "info", "  🧠 正在进行深度评估与破局计划生成...")
                    

                    full_jd_info = (
                        f"【岗位基本信息】\n"
                        f"薪资范围: {salary} | 工作城市: {city} | 经验要求: {experience} | 学历要求: {education}\n\n"
                        f"【岗位详情】\n{jd_text}"
                    )

                    first_stage_scores = {
                        "综合评级": feishu_api.extract_feishu_text(fields.get("综合评级 (A-F)", "")),
                        "核心-角色匹配": fields.get("核心-角色匹配"),
                        "高权-薪资契合": fields.get("高权-薪资契合"),
                        "AI评估详情": feishu_api.extract_feishu_text(fields.get("AI评估详情", "")),
                    }

                    import time
                    start_time = time.time()

                    try:
                        deep_result, deep_usage = await asyncio.to_thread(
                            deep_evaluate_resume,
                            resume_text,
                            full_jd_info,
                            first_stage_scores
                        )
                    except ValueError:
                        await send_sse_msg(queue, "error", "  ⚠️ AI 返回格式异常，已触发自动纠偏或跳过", job_id=job_id)
                        task_status[task_id]["completed"] = index
                        _sent_terminal = True
                        continue

                    elapsed_time = time.time() - start_time
                    pt = deep_usage.get("prompt_tokens", 0)
                    ct = deep_usage.get("completion_tokens", 0)
                    tt = deep_usage.get("total_tokens", 0)

                    await send_sse_msg(queue, "info", f"📈 Token 消耗 → 提示: {pt} / 补全: {ct}", job_id=job_id)
                    await send_sse_msg(queue, "info", f"  ✓ 深度评估完成 (耗时 {elapsed_time:.1f}s)", usage={"prompt": pt, "completion": ct, "total": tt})

                    extracted = deep_result.get("extracted_skills", [])
                    extracted_str = "、".join(extracted) if isinstance(extracted, list) else str(extracted)

                    fields_to_update = {
                        "理想画像与能力信号": str(deep_result.get("dream_picture", "") or deep_result.get("理想画像与能力信号", "")),
                        "核心能力词典": str(deep_result.get("ats_ability_analysis", "") or deep_result.get("核心能力词典", "")),
                        "高杠杆匹配点": str(deep_result.get("strong_fit_assessment", "") or deep_result.get("高杠杆匹配点", "")),
                        "致命硬伤与毒点": str(deep_result.get("risk_red_flags", "") or deep_result.get("致命硬伤与毒点", "")),
                        "破局行动计划": str(deep_result.get("deep_action_plan", "") or deep_result.get("破局行动计划", "")),
                        "跟进状态": "已完成深度评估",
                    }
                    if extracted_str:
                        fields_to_update["核心能力词典"] = (
                            f"【核心技能词条】{extracted_str}\n\n"
                            + fields_to_update["核心能力词典"]
                        )

                    await send_sse_msg(queue, "info", "  ☁️ 正在将深度评估结果回写到飞书...")

                    update_success = await asyncio.to_thread(
                        feishu_api.update_feishu_record,
                        record_id,
                        fields_to_update,
                        table_id
                    )

                    if update_success:
                        job_updates = {"followStatus": "已完成深度评估"}
                        await send_sse_msg(queue, "info", "  ✅ 深度评估完成，状态已更新为: 已完成深度评估")
                    else:
                        await send_sse_msg(queue, "warning", "  ⚠️ 深度评估完成，但回写飞书失败")

                elif task_type == "deliver":
                    # 🌟 如果有定时参数，将时间写入飞书并设状态为"待投递"，由 scheduler 接管
                    if scheduled_at:
                        from datetime import timezone, timedelta
                        try:
                            dt = datetime.strptime(scheduled_at, "%Y-%m-%d %H:%M:%S")
                        except ValueError:
                            try:
                                dt = datetime.strptime(scheduled_at, "%Y-%m-%d %H:%M")
                            except ValueError:
                                dt = None
                        if not dt:
                            raise Exception(f"无效的时间格式: {scheduled_at}")
                        # 🌟 核心修复：强制将提取出的时间认定为北京时间 (UTC+8)
                        tz_beijing = timezone(timedelta(hours=8))
                        dt = dt.replace(tzinfo=tz_beijing)
                        # 带有 tzinfo 的 dt 转 timestamp 时会准确换算为全球统一的绝对时间戳
                        timestamp_ms = int(dt.timestamp() * 1000)
                        ok = feishu_api.update_feishu_record(record_id, {
                            "定时投递时间": timestamp_ms,
                            "跟进状态": "待投递"
                        })
                        if not ok:
                            raise Exception(f"双保险写入失败：无法将 {record_id} 的定时投递信息写入飞书，请重试")
                        job_updates = {"followStatus": "待投递"}
                        await send_sse_msg(queue, "info", f"  ⏰ 定时投递已登记：{scheduled_at}，调度器将在到达时间后自动执行")
                        continue

                    await send_sse_msg(queue, "info", f"  🚀 正在准备自动投递物料...")

                    pdf_attachments = fields.get("PDF备份", [])
                    file_token = ""
                    pdf_name = "专属定制简历"
                    if pdf_attachments and isinstance(pdf_attachments, list) and len(pdf_attachments) > 0:
                        file_token = pdf_attachments[0].get("file_token", "")
                        pdf_name = pdf_attachments[0].get("name", "专属定制简历").replace(".pdf", "")

                    # 提取图片版简历（供 BOSS 投递引擎使用）
                    image_attachments = fields.get("图片保存", []) or []
                    image_items = []
                    if isinstance(image_attachments, list):
                        for att in image_attachments:
                            token = att.get("file_token", "")
                            if token:
                                image_items.append({"token": token, "name": att.get("name", "image.jpg")})

                    job_link_obj = fields.get("岗位链接", {})
                    job_url = job_link_obj.get("link", "") if isinstance(job_link_obj, dict) else str(job_link_obj)
                    greeting = feishu_api.extract_feishu_text(fields.get("打招呼语", ""))

                    if not (file_token or image_items) or not greeting or not job_url:
                        raise Exception("❌ 数据不全：缺少 [PDF备份或图片简历 / 打招呼语 / 岗位链接]，无法执行投递")

                    job_data = {
                        "record_id": record_id,
                        "job_url": job_url,
                        "file_token": file_token,
                        "pdf_name": pdf_name,
                        "greeting": greeting,
                        "image_items": image_items
                    }

                    platform_lower = platform.lower()
                    if "猎聘" in platform:
                        await send_sse_msg(queue, "info", "  🎯 路由匹配成功：调用 [猎聘] 全自动投递引擎...")
                        # 🌟 修复：延迟加载猎聘投递引擎，避免后端启动时误弹浏览器
                        liepin_dir = str(BASE_DIR.parent / "liepin_scraper")
                        if liepin_dir not in sys.path:
                            sys.path.insert(0, liepin_dir)
                        import liepin_auto_delivery
                        
                        success = await asyncio.to_thread(liepin_auto_delivery.deliver_job, job_data)
                        if success:
                            job_updates = {"followStatus": "已投递"}
                            await send_sse_msg(queue, "info", "  ✅ 投递成功，状态已更新为: 已投递")
                        else:
                            raise Exception("投递引擎执行失败，请检查控制台日志了解详情")
                    elif "boss" in platform_lower:
                        await send_sse_msg(queue, "info", "  🎯 路由匹配成功：调用 [BOSS直聘] 全自动投递引擎...")
                        boss_dir = str(BASE_DIR.parent / "boss_scraper")
                        if boss_dir not in sys.path:
                            sys.path.insert(0, boss_dir)
                        import boss_auto_delivery
                        success = await asyncio.to_thread(boss_auto_delivery.deliver_job, job_data)
                        if success:
                            job_updates = {"followStatus": "已投递"}
                            await send_sse_msg(queue, "info", "  ✅ 投递成功，状态已更新为: 已投递")
                        else:
                            raise Exception("BOSS 投递引擎执行失败，详情请看「自动投递失败日志」")
                    elif "51job" in platform_lower or "前程无忧" in platform:
                        await send_sse_msg(queue, "warning", "  ⚠️ [51job] 投递引擎暂未接入，已跳过")
                    else:
                        await send_sse_msg(queue, "warning", f"  ⚠️ 未知平台 [{platform}]，无法自动投递，已跳过")

                task_status[task_id]["completed"] = index
                await send_sse_msg(queue, "info", f"✅ 岗位 {job_id} 流程已完全结束，正在同步 UI...")
                await send_sse_msg(queue, "success", f"✓ {job_id} 处理完成", job_id=job_id, job_updates=job_updates)
                _sent_terminal = True

            except Exception as e:
                print(f"🚨 [处理过程发生致命异常]: {str(e)}")
                await send_sse_msg(queue, "error", f"  ✗ {job_id} 处理失败: {str(e)}", job_id=job_id)
                _sent_terminal = True
            finally:
                if not _sent_terminal:
                    await send_sse_msg(queue, "error", f"  ✗ {job_id} 处理意外中止", job_id=job_id)
        
        task_status[task_id]["status"] = "completed"
        task_status[task_id]["completed_at"] = datetime.now().isoformat()
        await send_sse_msg(queue, "complete", f"🎉 批量任务全部完成！共处理 {len(job_ids)} 个岗位")
        
    except Exception as e:
        task_status[task_id]["status"] = "failed"
        task_status[task_id]["error"] = str(e)
        
        import traceback
        print(f"\n🚨 [致命错误] 批量评估任务彻底崩溃: {str(e)}")
        traceback.print_exc()
        
        await send_sse_msg(queue, "error", f"❌ 批量任务执行失败: {str(e)}")
    finally:
        GLOBAL_TASK_STATE["is_processing"] = False  # 🔓 无论成功、失败还是异常，务必释放锁
        await queue.put('data: {"type": "end"}\n\n')


@app.post("/api/tasks/batch-process")
async def batch_process(payload: BatchTaskRequest, background_tasks: BackgroundTasks) -> Dict[str, Any]:
    print(f"\n👉 [后端] 收到批量任务请求，类型: {payload.task_type}，数量: {len(payload.job_ids)}")
    if payload.task_type not in ["evaluate", "rewrite", "deep_rewrite", "deep_evaluate", "deliver"]:
        raise HTTPException(
            status_code=400,
            detail="task_type 必须是 'evaluate'、'rewrite'、'deep_rewrite'、'deep_evaluate' 或 'deliver'"
        )
    
    if not payload.job_ids:
        raise HTTPException(
            status_code=400,
            detail="job_ids 不能为空"
        )
    
    task_id = str(uuid.uuid4())
    task_queues[task_id] = asyncio.Queue()
    
    task_status[task_id] = {
        "status": "pending",
        "task_type": payload.task_type,
        "total": len(payload.job_ids),
        "completed": 0,
        "created_at": datetime.now().isoformat()
    }
    
    background_tasks.add_task(
        run_batch_ai_task,
        task_id=task_id,
        task_type=payload.task_type,
        job_ids=payload.job_ids,
        queue=task_queues[task_id],
        scheduled_at=payload.scheduled_at
    )
    
    return {
        "status": "started",
        "task_id": task_id,
        "message": f"批量任务已启动，共 {len(payload.job_ids)} 个岗位"
    }

# ----------------------------------------------------------------
# 🌟 新增：简历视觉智能解析接口 (PDF/Word -> Vision -> Markdown)
# ----------------------------------------------------------------

@app.post("/api/strategy/upload_resume_vision")
async def upload_resume_vision(file: UploadFile = File(...)):
    import io
    from datetime import datetime
    import base64
    
    # 临时文件路径（解析完会删除）
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    temp_pdf = BASE_DIR / f"temp_upload_{timestamp}.pdf"
    temp_docx = BASE_DIR / f"temp_upload_{timestamp}.docx"
    
    try:
        content = await file.read()
        image_base64_list = []

        # --- 步骤 1：统一转换为图片 ---
        if file.filename.endswith(".pdf"):
            with open(temp_pdf, "wb") as f: f.write(content)
            pages = convert_from_path(str(temp_pdf), dpi=200)
            for page in pages:
                buf = io.BytesIO()
                page.save(buf, format="JPEG")
                image_base64_list.append(base64.b64encode(buf.getvalue()).decode('utf-8'))
                
        elif file.filename.endswith(".docx"):
            with open(temp_docx, "wb") as f: f.write(content)
            # 复用你已有的 docx2pdf 逻辑
            convert(str(temp_docx), str(temp_pdf))
            pages = convert_from_path(str(temp_pdf), dpi=200)
            for page in pages:
                buf = io.BytesIO()
                page.save(buf, format="JPEG")
                image_base64_list.append(base64.b64encode(buf.getvalue()).decode('utf-8'))
        
        if not image_base64_list:
            return {"status": "error", "detail": "未能生成有效预览图"}

        # --- 步骤 2：呼叫视觉大模型 ---
        system_prompt = """你是一个顶级的简历重构专家。
请观察图片中的简历内容，将其还原为结构清晰的文本。
【强制排版规范】：
1. 仅保留一级大标题（如个人总结、项目经历等），并且必须且仅能使用单个 # 开头（例：#项目经历）。
2. 模块内的所有正文（包括具体的公司名、项目名、工作内容等）直接换行排列即可，**绝对不要**使用 ##、### 等任何多级标题或 Markdown 标记符号。
3. 保持简历内容的原始完整性，不要自行精简或扩写。
4. 直接输出结果文本，不要包含 ```markdown 等代码块包裹符号。"""

        user_content = [{"type": "text", "text": "请将这张/这些简历截图解析为指定的 Markdown 格式："}]
        for b64 in image_base64_list:
            user_content.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{b64}"}
            })

        print(f">>> 正在使用视觉模型解析简历，共 {len(image_base64_list)} 页...")
        # 🌟 复用你 main.py 里的 VISION_LLM_MODEL
        response = client.chat.completions.create(
            model=VISION_LLM_MODEL, 
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ],
            temperature=0.1
        )
        
        markdown_result = (response.choices[0].message.content or "").strip()
        return {"status": "success", "text": markdown_result}

    # 🚨 这里就是刚才不小心被删掉的“尾巴”，现在补回来了！
    except Exception as e:
        print(f"❌ 简历视觉解析失败: {str(e)}")
        return {"status": "error", "detail": str(e)}
    finally:
        # 清理垃圾
        for p in [temp_pdf, temp_docx]:
            if p.exists(): p.unlink()

# ==================== 🌟 ChatOps 自动化任务编排 (多平台/防风控) ====================

class ChatCommandRequest(BaseModel):
    command: str
    task_id: Optional[str] = None  # 🌟 前端可以传入自己生成的 task_id

# 1. 平台单页数据量基准配置
PLATFORM_CONFIG = {
    "boss": {"per_page": 15, "script_name": "boss_collector.py", "dir_name": "boss_scraper"},
    "51job": {"per_page": 20, "script_name": "51job_collector.py", "dir_name": "51job_scraper"},
    "liepin": {"per_page": 30, "script_name": "liepin_crawler.py", "dir_name": "liepin_scraper"},
    "zhaopin": {"per_page": 30, "script_name": "data_collector.py", "dir_name": "zhilian_scraper"},
    "智联招聘": {"per_page": 30, "script_name": "data_collector.py", "dir_name": "zhilian_scraper"} 
}

# 🌟 全局控制字典：用于拦截和强杀正在运行的爬虫子进程
active_processes: Dict[str, asyncio.subprocess.Process] = {}
cancel_events: Dict[str, asyncio.Event] = {}

def calculate_pages(target_count: int, platform: str) -> int:
    plat_key = platform.lower()
    if plat_key == "boss直聘": plat_key = "boss"
    if plat_key == "前程无忧": plat_key = "51job"
    
    per_page = PLATFORM_CONFIG.get(plat_key, {}).get("per_page", 20)
    pages = int((target_count / per_page) * 1.2)
    return max(1, pages)

async def parse_chat_intent(command: str) -> dict:
    if not client:
        raise HTTPException(status_code=500, detail="AI 服务未配置，无法解析指令")

    # 🌟 全能加强版 LLM Prompt：保留原始 Schema 细节，新增清洗与同步能力
    system_prompt = """你是一个爬虫任务编排与数据库分析助手。请从用户指令中提取参数。
必须严格输出纯净 JSON，绝不能包含 Markdown 代码块标记（如 ```json）。

【提取规则】：
0. "action": 字符串。必填字段，用于识别用户意图：
   - "scrape": 抓取、爬取岗位数据
   - "evaluate": 评估、打分简历或岗位
   - "clean": 执行 Python 结构化硬清洗 (调用 3_sqlite_smart_filter.py)
   - "push": 将清洗通过的岗位同步到飞书 (调用 sync_to_feishu.py)
   - "query": 查询数据状态、统计数量、看数据库（如"有多少没同步"、"今天抓了多少"）

1. 当 action 为 "query" 时，你必须生成一条合法的 SQLite 查询语句，并放入 "sql" 字段中。
   - 数据表名：raw_jobs
   - 核心字段：job_link, job_title, company_name, city, salary, platform, process_status, is_synced(0或NULL为未同步, 1为已同步), crawl_time(DATETIME)
   - ⚠️ 注意时间过滤：crawl_time 是本地时间，查询今天的数据请用：date(crawl_time) = date('now', 'localtime')
   - 示例1："今天抓了多少" -> "SELECT platform, count(*) as 数量 FROM raw_jobs WHERE date(crawl_time) = date('now', 'localtime') GROUP BY platform"
   - 示例2："有多少没push到飞书" -> "SELECT count(*) as 未同步数量 FROM raw_jobs WHERE is_synced = 0 OR is_synced IS NULL"

2. "platforms": 数组。如 ["boss", "liepin", "51job"]。
3. "specific_page": 整数。明确指定抓取"第X页"时提取。
4. "pages": 整数。总共"抓取X页"。
5. "target_count": 整数。抓取或评估的"个数"。若未提，默认 100。用户确认推送时，可提取具体的推送数量。
6. "keyword" / "city" / "salary": 字符串，搜索条件，"salary" 必须是大写 K 结尾（如 15-20K）。
7. "start_page": 整数。从第X页开始。
"""
    try:
        response = await asyncio.to_thread(
            client.chat.completions.create,
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"用户指令：{command}"}
            ],
            temperature=0.1
        )
        # 接收并清理大模型返回的结果
        content = (response.choices[0].message.content or "").strip()
        
        if content.startswith("```json"): 
            content = content[7:]
        elif content.startswith("```"): 
            content = content[3:]
        if content.endswith("```"): 
            content = content[:-3]
        
        import json
        return json.loads(content.strip())
    # 🌟 修复点：必须补上这个 except 块，否则会报 SyntaxError
    except Exception as e:
        print(f"⚠️ LLM 意图解析失败，使用默认降级参数: {e}")
        return {"action": "scrape", "platforms": ["boss"], "target_count": 100}

async def run_sequential_chatops_scheduler(task_id: str, platforms: list, target_count: int, target_pages: int, specific_page: int, keyword: str, city: str, salary: str, queue: asyncio.Queue, start_page: int = 1):
    cancel_event = asyncio.Event()
    cancel_events[task_id] = cancel_event
    
    try:
        mode_msg = f"第 {specific_page} 页" if specific_page is not None and specific_page > 0 else (f"{target_pages} 页" if target_pages else f"各 {target_count} 条")
        await queue.put(f'data: {{"type": "info", "message": "🚦 调度启动：识别到 {len(platforms)} 个平台，执行模式：{mode_msg}数据。"}}\n\n')

        for plat in platforms:
            # 🌟 动态获取当前平台的爬虫目录
            platform_dir_name = PLATFORM_CONFIG.get(plat, {}).get('dir_name', 'boss_scraper')
            spider_dir = BASE_DIR.parent / platform_dir_name
            
            # 🌟 第一步：如果是 boss，执行登录并实时吐出日志
            if plat == "boss":
                await queue.put(f'data: {{"type": "info", "message": "🔐 正在自动提取 Edge 浏览器 Cookie 刷新登录态..."}}\n\n')
                await queue.put(f'data: {{"type": "progress", "message": "  > boss login --cookie-source edge"}}\n\n')
                
                login_process = await asyncio.create_subprocess_exec(
                    "boss", "login", "--cookie-source", "edge",
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE,
                    cwd=str(spider_dir)
                )

                async def read_login_stream(stream, is_error=False):
                    while True:
                        line = await stream.readline()
                        if not line: break
                        text = line.decode('utf-8', errors='replace').strip()
                        if text:
                            safe_text = text.replace('"', '\\"').replace('\n', ' ')
                            msg_type = "warning" if is_error else "info"
                            await queue.put(f'data: {{"type": "{msg_type}", "message": "  [Cookie刷新] {safe_text}"}}\n\n')

                await asyncio.gather(
                    read_login_stream(login_process.stdout), 
                    read_login_stream(login_process.stderr, True)
                )
                await login_process.wait()
                await queue.put(f'data: {{"type": "success", "message": "✅ 登录态刷新完毕，准备拉起采集矩阵！"}}\n\n')

            # 🌟 第二步：精准单页 vs 动态目标循环
            if specific_page is not None and specific_page > 0:
                # 精准模式：只抓指定的那一页，原样执行
                page = specific_page
                await queue.put(f'data: {{"type": "info", "message": "📍 精准模式：锁定抓取 {plat} 的第 {specific_page} 页"}}\n\n')
                if cancel_event.is_set(): raise Exception("任务已被用户手动终止")

                await queue.put(f'data: {{"type": "progress", "message": "🚀 正在拉起 {plat} 爬虫，执行指令..."}}\n\n')

                script_name = PLATFORM_CONFIG.get(plat, {}).get('script_name', 'boss_collector.py')
                display_cmd = f"python {script_name} -p {page}"
                await queue.put(f'data: {{"type": "progress", "message": "  > {display_cmd}"}}\n\n')

                script_path = str(spider_dir / script_name)
                cmd = ["python", "-u", script_path, "--platform", plat, "-p", str(page)]

                process = await asyncio.create_subprocess_exec(
                    *cmd,
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE,
                    cwd=str(spider_dir)
                )
                active_processes[task_id] = process

                async def read_stream_single(stream, msg_type="info"):
                    while True:
                        line = await stream.readline()
                        if not line: break
                        text = line.decode('utf-8', errors='replace').strip()
                        if text:
                            safe_text = text.replace('"', '\\"').replace('\n', ' ')
                            await queue.put(f'data: {{"type": "{msg_type}", "message": "  [{plat}-P{page}] {safe_text}"}}\n\n')
                            await asyncio.sleep(0.01)  # 主动让出 Event Loop，防止高频日志独占

                await asyncio.gather(read_stream_single(process.stdout, "info"), read_stream_single(process.stderr, "warning"))
                await process.wait()

            else:
                # 🌟 动态目标循环模式：按实际入库量决定是否继续翻页
                tracker = {"inserted": 0}
                page = start_page
                max_pages = start_page + 29

                # ---- A. 配置决策：有临时关键词用临时，否则拉飞书默认启用配置 ----
                config_record_id = None
                base_count = 0
                final_keyword = keyword
                final_city = city
                final_salary = salary

                if not keyword:
                    try:
                        feishu_config = await asyncio.to_thread(feishu_api.get_enabled_search_config)
                        if feishu_config:
                            config_record_id = feishu_config["record_id"]
                            base_count = feishu_config["抓取数量"]
                            final_keyword = feishu_config["岗位Title"]
                            final_city = feishu_config["城市"]
                            final_salary = feishu_config["薪资"]
                            display_city = final_city or "全国"
                            await queue.put(f'data: {{"type": "info", "message": "📊 使用飞书默认配置：[{display_city}] [{final_keyword}]"}}\n\n')
                        else:
                            await queue.put(f'data: {{"type": "warning", "message": "⚠️ 飞书搜索配置表无启用记录，使用空参数继续..."}}\n\n')
                    except Exception as cfg_e:
                        safe_cfg_err = str(cfg_e).replace('"', '\\"')
                        await queue.put(f'data: {{"type": "warning", "message": "⚠️ 读取飞书配置失败: {safe_cfg_err}，继续执行..."}}\n\n')
                else:
                    display_city = final_city or "全国"
                    await queue.put(f'data: {{"type": "info", "message": "🎯 收到临时指令，使用条件：[{display_city}] [{final_keyword}]"}}\n\n')

                await queue.put(f'data: {{"type": "info", "message": "🎯 动态模式：目标入库 {target_count} 条，从第 1 页开始，最多抓 {max_pages} 页"}}\n\n')

                while tracker["inserted"] < target_count and page <= max_pages:
                    if cancel_event.is_set(): raise Exception("任务已被用户手动终止")

                    await queue.put(f'data: {{"type": "progress", "message": "🚀 正在拉起 {plat} 爬虫，执行指令..."}}\n\n')

                    script_name = PLATFORM_CONFIG.get(plat, {}).get('script_name', 'boss_collector.py')
                    display_cmd = f"python {script_name} -p {page}"
                    await queue.put(f'data: {{"type": "progress", "message": "  > {display_cmd}"}}\n\n')

                    script_path = str(spider_dir / script_name)
                    cmd = ["python", "-u", script_path, "--platform", plat, "-p", str(page)]
                    if final_keyword: cmd.extend(["--keyword", final_keyword])
                    if final_city: cmd.extend(["--city", final_city])
                    if final_salary: cmd.extend(["--salary", final_salary])

                    process = await asyncio.create_subprocess_exec(
                        *cmd,
                        stdout=asyncio.subprocess.PIPE,
                        stderr=asyncio.subprocess.PIPE,
                        cwd=str(spider_dir)
                    )
                    active_processes[task_id] = process

                    async def read_stream(stream, msg_type="info", _page=page, _tracker=tracker):
                        while True:
                            line = await stream.readline()
                            if not line: break
                            text = line.decode('utf-8', errors='replace').strip()
                            if text:
                                safe_text = text.replace('"', '\\"').replace('\n', ' ')
                                await queue.put(f'data: {{"type": "{msg_type}", "message": "  [{plat}-P{_page}] {safe_text}"}}\n\n')
                                await asyncio.sleep(0.01)  # 主动让出 Event Loop，防止高频日志独占
                                match = re.search(r"新增入库\s*(\d+)\s*个", text)
                                if match:
                                    _tracker["inserted"] += int(match.group(1))

                    await asyncio.gather(read_stream(process.stdout, "info"), read_stream(process.stderr, "warning"))
                    await process.wait()

                    page += 1

                    # 尚未达标且还有剩余页数配额，执行防风控休眠后继续
                    if tracker["inserted"] < target_count and not cancel_event.is_set() and page <= max_pages:
                        inserted_so_far = tracker["inserted"]
                        await queue.put(f'data: {{"type": "info", "message": "📈 目标 {target_count} 个，当前已入库 {inserted_so_far} 个，自动开启第 {page} 页抓取..."}}\n\n')
                        await queue.put(f'data: {{"type": "warning", "message": "🛡️ 防风控保护：休眠 10 分钟后继续..."}}\n\n')
                        wait_seconds = 600
                        elapsed = 0
                        while elapsed < wait_seconds:
                            if cancel_event.is_set(): raise Exception("任务已在休眠期间手动终止")
                            
                            remaining = wait_seconds - elapsed
                            
                            # 每60秒播报一次倒计时
                            if remaining % 60 == 0 and remaining > 0:
                                minutes = remaining // 60
                                if minutes > 1:
                                    await queue.put(f'data: {{"type": "info", "message": "⏱️ 休眠倒计时：还剩 {minutes} 分钟..."}}\n\n')
                                elif minutes == 1:
                                    await queue.put(f'data: {{"type": "info", "message": "⏱️ 休眠倒计时：还剩 1 分钟，准备唤醒..."}}\n\n')
                            
                            # 每10秒发送心跳，防止SSE连接超时断开
                            if elapsed % 10 == 0:
                                await queue.put(f'data: {{"type": "heartbeat", "message": ""}}\n\n')
                            
                            await asyncio.sleep(1)
                            elapsed += 1

                # ---- B. 抓取完成后回写飞书搜索配置表 ----
                try:
                    final_inserted = tracker["inserted"]
                    if final_inserted > 0:
                        # 先播报进度提示，制造无感体验
                        await queue.put(f'data: {{"type": "info", "message": "⏳ 正在将本次抓取的 {final_inserted} 条数据同步至飞书配置中心..."}}\n\n')
                        
                        # 使用异步线程池包裹同步 HTTP 调用，避免阻塞事件循环
                        if config_record_id:
                            await asyncio.to_thread(feishu_api.update_search_config_count, config_record_id, base_count, final_inserted)
                        else:
                            await asyncio.to_thread(feishu_api.create_search_config, final_keyword or "", final_city or "", final_salary or "", final_inserted)
                        
                        # 回写完成后再发送成功消息
                        await queue.put(f'data: {{"type": "success", "message": "✅ 抓取数据已成功回写至飞书搜索配置表，本次共入库 {final_inserted} 条"}}\n\n')
                except Exception as wb_e:
                    safe_wb_err = str(wb_e).replace('"', '\\"')
                    await queue.put(f'data: {{"type": "warning", "message": "⚠️ 飞书回写失败（不影响数据入库）: {safe_wb_err}"}}\n\n')

        await queue.put(f'data: {{"type": "success", "message": "✨ 所有抓取任务已安全执行完毕！"}}\n\n')

    except Exception as e:
        await queue.put(f'data: {{"type": "error", "message": "🛑 {str(e)}"}}\n\n')
    finally:
        await queue.put('data: {"type": "end"}\n\n')
        if task_id in active_processes: del active_processes[task_id]
        if task_id in cancel_events: del cancel_events[task_id]

async def run_chatops_evaluator(task_id: str, target_count: int, platforms: list, queue: asyncio.Queue):
    """
    ChatOps 批量评估调度器：从飞书拉取未评估岗位，逐一调用 AI 评估并回写
    """
    cancel_event = asyncio.Event()
    cancel_events[task_id] = cancel_event
    
    try:
        await send_sse_msg(queue, "info", f"🧠 收到评估指令，正在从飞书总表拉取最新 {target_count} 个未评估岗位...")
        
        # 步骤 1: 从飞书拉取未评估岗位（综合评级为空或 AI总分为0）
        all_records = await asyncio.to_thread(
            feishu_api.get_feishu_records,
            FEISHU_TABLE_ID_JOBS,
            filter_conditions=None
        )
        
        if not all_records:
            await send_sse_msg(queue, "warning", "⚠️ 飞书总表中暂无数据，请先抓取岗位")
            return
        
        # 筛选未评估的岗位（综合评级为空）
        unevaluated_jobs = []
        for record in all_records:
            fields = record.get("fields", {})
            grade = feishu_api.extract_feishu_text(fields.get("综合评级", ""))
            
            # 如果有平台筛选，只评估指定平台的岗位
            platform = feishu_api.extract_feishu_text(fields.get("招聘平台", "") or fields.get("数据来源", ""))
            if platforms and platform:
                platform_lower = platform.lower()
                matched = False
                for p in platforms:
                    if p in platform_lower or platform_lower in p:
                        matched = True
                        break
                if not matched:
                    continue
            
            if not grade or grade.strip() == "":
                unevaluated_jobs.append({
                    "record_id": record.get("record_id"),
                    "table_id": FEISHU_TABLE_ID_JOBS,
                    "platform": platform or "未知平台",
                    "company": feishu_api.extract_feishu_text(fields.get("公司名称", "未知公司")),
                    "job_title": feishu_api.extract_feishu_text(fields.get("岗位名称", "未知岗位")),
                    "jd_text": feishu_api.extract_feishu_text(fields.get("岗位详情", "")),
                    "salary": feishu_api.extract_feishu_text(fields.get("薪资", "")),
                    "city": feishu_api.extract_feishu_text(fields.get("城市", "")),
                    "experience": feishu_api.extract_feishu_text(fields.get("经验要求", "")),
                    "education": feishu_api.extract_feishu_text(fields.get("学历要求", "")),
                })
        
        if not unevaluated_jobs:
            await send_sse_msg(queue, "info", "✅ 所有岗位均已评估完毕，无需重复评估")
            return
        
        # 限制数量
        jobs_to_eval = unevaluated_jobs[:target_count]
        total = len(jobs_to_eval)
        
        await send_sse_msg(queue, "info", f"📊 找到 {len(unevaluated_jobs)} 个未评估岗位，本次将评估前 {total} 个")
        
        # 步骤 2: 加载简历
        await send_sse_msg(queue, "info", "📄 正在加载云端简历底稿...")
        resume_text = await asyncio.to_thread(load_resume_evaluator)
        
        if not resume_text:
            await send_sse_msg(queue, "error", "❌ 未找到启用的简历，请在飞书配置中心设置")
            return
        
        await send_sse_msg(queue, "success", "✅ 简历加载成功，开始批量评估...")

        # 步骤 2.5: 预加载求职偏好（全局仅加载一次）
        await send_sse_msg(queue, "info", "⚙️ 正在加载求职偏好与底线配置...")
        preferences_text = await asyncio.to_thread(feishu_api.get_my_preferences)
        if preferences_text:
            await send_sse_msg(queue, "success", "✅ 偏好配置加载成功，AI 将执行一票否决校验")
        else:
            await send_sse_msg(queue, "warning", "⚠️ 未找到求职偏好配置，将使用通用评估模式")

        # 步骤 3: 逐一评估
        success_count = 0
        fail_count = 0
        
        for idx, job in enumerate(jobs_to_eval, 1):
            # 检查终止信号
            if cancel_event.is_set():
                await send_sse_msg(queue, "warning", "⏸️ 收到终止信号，评估任务已中断")
                break
            
            company = job["company"]
            job_title = job["job_title"]
            platform = job["platform"]
            
            await send_sse_msg(queue, "progress", f"⏳ 正在评估 {idx}/{total}: [{platform}] {company} - {job_title}...")

            # 背调公司情报
            await send_sse_msg(queue, "info", f"  🔍 正在背调【{company}】公司情报...")
            company_intel = await research_company_serper(company)
            await send_sse_msg(queue, "info", f"  📡 情报获取完毕: {company_intel[:60]}...")
            
            try:
                # 调用评估函数（传入公司情报 + 个人偏好）
                result = await asyncio.to_thread(
                    evaluate_single_job,
                    job,
                    resume_text,
                    company_intel,
                    preferences_text,
                )
                
                if result["success"]:
                    grade = result.get("grade", "?")
                    ai_score = result.get("ai_score", 0)
                    usage = result.get("usage", {})
                    pt = usage.get("prompt_tokens", 0)
                    ct = usage.get("completion_tokens", 0)
                    tt = usage.get("total_tokens", 0)
                    
                    # 回写飞书
                    update_success = await asyncio.to_thread(
                        feishu_api.update_feishu_record,
                        job["record_id"],
                        result["update_data"],
                        FEISHU_TABLE_ID_JOBS
                    )
                    
                    if update_success:
                        await send_sse_msg(queue, "success", f"  ✅ 评估完成！综合评级: {grade}，AI总分: {ai_score}分", usage={"prompt": pt, "completion": ct, "total": tt})
                        success_count += 1
                    else:
                        await send_sse_msg(queue, "warning", "  ⚠️ 评估完成但回写飞书失败")
                        fail_count += 1
                else:
                    error_msg = result.get("error", "未知错误")
                    await send_sse_msg(queue, "warning", f"  ⚠️ 评估失败: {error_msg}")
                    fail_count += 1
                    
            except Exception as e:
                await send_sse_msg(queue, "warning", f"  ⚠️ 评估异常: {str(e)}")
                fail_count += 1
            
            # 短暂延时避免过载
            await asyncio.sleep(0.5)
        
        # 步骤 4: 总结
        await send_sse_msg(queue, "success", f"🎉 批量评估完成！成功 {success_count} 个，失败 {fail_count} 个，数据已全部回写")
        
    except Exception as e:
        await send_sse_msg(queue, "error", f"🛑 评估调度异常: {str(e)}")
    finally:
        await send_sse_msg(queue, "end", "")
        if task_id in cancel_events: del cancel_events[task_id]

async def research_company_serper(company_name: str) -> str:
    """
    调用 Serper.dev 搜索公司融资/规模/产品情报，返回结构化背景文本（约 300 字）。
    若 Key 为空或请求失败则优雅降级。
    """
    if "某" in company_name or company_name == "未知公司":
        return "⚠️ 匿名或未知公司，跳过外部背调"

    # 🔍 调试：打印 Key 加载状态
    raw_key = os.getenv("SERPER_API_KEY")
    if raw_key:
        print(f"   🔑 [Serper] SERPER_API_KEY 已加载，前4位: {raw_key[:4]}****")
    else:
        print(f"   ❌ [Serper] SERPER_API_KEY 未加载（os.getenv 返回 None）")

    if not SERPER_API_KEY:
        return "⚠️ 情报获取失败，降级评估（未配置 SERPER_API_KEY）"

    def _fetch():
        try:
            url = "https://google.serper.dev/search"
            headers = {"X-API-KEY": SERPER_API_KEY, "Content-Type": "application/json"}
            payload = {
                "q": f"{company_name} 公司介绍 核心业务",
                "gl": "cn",
                "hl": "zh-cn",
                "num": 5,
            }
            resp = requests.post(url, headers=headers, json=payload, timeout=10)
            print(f"   📡 [Serper] HTTP {resp.status_code}，响应长度: {len(resp.text)}")
            data = resp.json()

            snippets = []
            for item in data.get("organic", [])[:5]:
                snippet = item.get("snippet", "").strip()
                if snippet:
                    snippets.append(snippet)

            if not snippets:
                print(f"   ⚠️ [Serper] 响应中无 organic snippets，原始 keys: {list(data.keys())}")
                return "⚠️ 情报获取失败，降级评估"

            intel = f"【{company_name} 公司情报】\n" + "\n".join(f"· {s}" for s in snippets[:3])
            return intel[:500]
        except Exception as e:
            import traceback
            print(f"   ❌ [Serper] 请求异常: {type(e).__name__}: {e}")
            traceback.print_exc()
            return f"⚠️ 情报获取失败，降级评估（{str(e)[:60]}）"

    return await asyncio.to_thread(_fetch)


# 1. 结构化清洗执行器
async def run_chatops_cleaner(task_id: str, queue: asyncio.Queue):
    try:
        await send_sse_msg(queue, "info", "🧹 正在拉起 Python 结构化硬清洗与打分引擎...")
        # 🌟 修复 1：指向新的 step1 脚本
        script_path = str(BASE_DIR.parent / "job_processor" / "step1_rule_filter.py")
        
        process = await asyncio.create_subprocess_exec(
            sys.executable, "-u", script_path,
            stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.STDOUT,
            cwd=str(BASE_DIR.parent)
        )
        
        while True:
            line = await process.stdout.readline()
            if not line: break
            await send_sse_msg(queue, "progress", f" > {line.decode().strip()}")
            
        await process.wait()

        # 🌟 修复 2：统计待推送数量（兼容新版打分机制及 is_synced 字段可能缺失的情况）
        def _get_pending_count():
            with sqlite3.connect(DB_PATH) as conn:
                cur = conn.cursor()
                try:
                    # 尝试包含 is_synced 的查询
                    cur.execute("SELECT count(*) FROM raw_jobs WHERE keywords_score > 20 AND (is_synced = 0 OR is_synced IS NULL)")
                except sqlite3.OperationalError:
                    # 若 step2 尚未执行过，is_synced 列可能还未创建，降级查询
                    cur.execute("SELECT count(*) FROM raw_jobs WHERE keywords_score > 20")
                return cur.fetchone()[0]

        count = await asyncio.to_thread(_get_pending_count)
        await send_sse_msg(queue, "success", f"✅ 清洗与打分完成！当前有 【 {count} 】 个高分岗位就绪。你可以指令我：'把这些岗位推送到飞书'。")
    except Exception as e:
        await send_sse_msg(queue, "error", f"❌ 清洗出错: {str(e)}")
    finally:
        await queue.put('data: {"type": "end"}\n\n')

# 2. 飞书同步执行器
async def run_chatops_pusher(task_id: str, queue: asyncio.Queue):
    try:
        await send_sse_msg(queue, "info", "🚀 收到确认！正在同步合格岗位至飞书...")
        # 🌟 修复 3：指向新的 step2 脚本
        script_path = str(BASE_DIR.parent / "job_processor" / "step2_sync_feishu.py")
        
        process = await asyncio.create_subprocess_exec(
            sys.executable, "-u", script_path,
            stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.STDOUT,
            cwd=str(BASE_DIR.parent)
        )
        
        push_count = 0
        while True:
            line = await process.stdout.readline()
            if not line: break
            text = line.decode().strip()
            await send_sse_msg(queue, "progress", f" > {text}")
            if "写入飞书成功" in text or "成功推送" in text:
                push_count += 1
                
        await process.wait()
        await send_sse_msg(queue, "success", f"🎉 同步完成！本次已成功推送 {push_count} 条数据到飞书。")
    except Exception as e:
        await send_sse_msg(queue, "error", f"❌ 同步失败: {str(e)}")
    finally:
        await queue.put('data: {"type": "end"}\n\n')

async def run_chatops_db_querier(task_id: str, sql: str, queue: asyncio.Queue):
    """
    ChatOps 数据库查询引擎：执行 LLM 生成的 SQL 并聚合结果输出
    """
    import sqlite3
    try:
        if not sql:
            await send_sse_msg(queue, "error", "❌ AI 未能生成有效的 SQL 查询语句")
            return
            
        await send_sse_msg(queue, "info", f"📊 正在执行查询...")
        await send_sse_msg(queue, "progress", f"SQL: {sql}")
        
        def _execute_query():
            with sqlite3.connect(DB_PATH) as conn:
                cursor = conn.cursor()
                cursor.execute(sql)
                rows = cursor.fetchall()
                # 提取列名
                cols = [description[0] for description in cursor.description]
                return rows, cols

        rows, cols = await asyncio.to_thread(_execute_query)
        
        if not rows:
            await send_sse_msg(queue, "warning", "⚠️ 查询结果为空。")
        else:
            # 🌟 核心优化 1：最大显示行数保护，防止几千条数据卡死前端
            limit = 30 
            display_rows = rows[:limit]
            
            # 🌟 核心优化 2：拼接成极度兼容的 Markdown 表格格式
            msg_lines = [
                f"**✅ 查询成功！(共 {len(rows)} 条记录)**",
                "",  # 🌟 关键修复：Markdown 语法要求表格上方必须有纯空行隔离
                "| " + " | ".join(cols) + " |",
                "| " + " | ".join(["---"] * len(cols)) + " |"
            ]
            
            for row in display_rows:
                # 把每行数据转为字符串，并处理空值
                row_str = " | ".join(str(item) if item is not None else "-" for item in row)
                msg_lines.append(f"| {row_str} |")
            
            if len(rows) > limit:
                msg_lines.append("")  # 🌟 关键修复：表格结束后也加纯空行断开
                msg_lines.append(f"*(⚠️ 数据过多，为保证渲染性能，仅展示前 {limit} 条)*")
            
            # 🌟 核心优化 3：将整个表格合并为一条消息，一次性发给前端
            final_message = "\n".join(msg_lines)
            await send_sse_msg(queue, "success", final_message)

    except Exception as e:
        await send_sse_msg(queue, "error", f"❌ 数据库查询异常: {str(e)}")
    finally:
        await queue.put('data: {"type": "end"}\n\n')

@app.post("/api/chat/command")
async def handle_chat_command(payload: ChatCommandRequest, background_tasks: BackgroundTasks) -> Dict[str, Any]:
    command_text = payload.command.strip()
    
    # 🌟 核心拦截逻辑：只要看到终止口令，直接切断不走 LLM
    if command_text in ["终止", "停止", "结束", "stop", "退出"]:
        for t_id, event in cancel_events.items():
            event.set() # 触发终止信号
            if t_id in active_processes:
                try: active_processes[t_id].terminate() # 强杀底层的 Python 进程
                except: pass
        return {"status": "success", "message": "已收到指令，正在强制停止所有后台进程..."}

    intent = await parse_chat_intent(command_text)
    action = intent.get("action", "scrape")
    platforms = intent.get("platforms", ["boss"])
    target_count = intent.get("target_count", 100)
    target_pages = intent.get("pages", 0)
    specific_page = intent.get("specific_page", 0)
    keyword = intent.get("keyword") or ""
    city = intent.get("city") or ""
    salary = intent.get("salary") or ""
    start_page = int(intent.get("start_page") or 1)

    # 🌟 优先使用前端传入的 task_id，如果没有则后端生成
    task_id = payload.task_id if payload.task_id else f"chatops_{uuid.uuid4().hex[:8]}"
    task_queues[task_id] = asyncio.Queue()
    
    if "task_status" in globals():
        task_type = "chat_cleaner" if action == "clean" else ("chat_evaluator" if action == "evaluate" else "chat_scraper")
        task_status[task_id] = {
            "status": "pending",
            "task_type": task_type,
            "created_at": datetime.now().isoformat()
        }
    
    # 🌟 根据 action 分流到不同的调度器
    if action == "clean":
        background_tasks.add_task(
            run_chatops_cleaner,
            task_id=task_id,
            queue=task_queues[task_id]
        )
    elif action == "push":
        # 🌟 新增：路由到同步执行器
        background_tasks.add_task(
            run_chatops_pusher,
            task_id=task_id,
            queue=task_queues[task_id]
        )    
    elif action == "query":
        # 🌟 新增：数据库查询路由
        background_tasks.add_task(
            run_chatops_db_querier,
            task_id=task_id,
            sql=intent.get("sql", ""),
            queue=task_queues[task_id]
        )
    elif action == "evaluate":
        background_tasks.add_task(
            run_chatops_evaluator,
            task_id=task_id,
            target_count=target_count,
            platforms=platforms,
            queue=task_queues[task_id]
        )
    else:
        background_tasks.add_task(
            run_sequential_chatops_scheduler,
            task_id=task_id,
            platforms=platforms,
            target_count=target_count,
            target_pages=target_pages,
            specific_page=specific_page,
            keyword=keyword,
            city=city,
            salary=salary,
            start_page=start_page,
            queue=task_queues[task_id]
        )
    
    return {
        "status": "started",
        "task_id": task_id,
        "planned_platforms": platforms,
        "target_per_platform": target_count
    }   

@app.get("/api/tasks/status")
def get_task_status():
    """短轮询接口：查询当前是否有后台任务正在执行"""
    return {"is_processing": GLOBAL_TASK_STATE.get("is_processing", False)}


@app.get("/api/tasks/logs")
async def task_logs(task_id: str):
    if task_id not in task_queues:
        raise HTTPException(
            status_code=404,
            detail=f"任务 {task_id} 不存在或已过期"
        )
    
    queue = task_queues[task_id]
    
    async def event_generator():
        try:
            yield f'data: {{"type": "connected", "message": "已连接到任务日志流", "task_id": "{task_id}"}}\n\n'
            
            while True:
                try:
                    # 🌟 缩短超时至 2s：更及时地发送心跳，防止 SSE 连接因空闲被代理/浏览器切断
                    message = await asyncio.wait_for(queue.get(), timeout=2.0)
                    if '"type": "end"' in message:
                        yield message
                        break
                    yield message
                except asyncio.TimeoutError:
                    # 2 秒内无新日志 → 发送心跳保持连接存活
                    yield f'data: {{"type": "heartbeat", "timestamp": "{datetime.now().isoformat()}"}}\n\n'
                except asyncio.CancelledError:
                    # 🌟 客户端断开连接时只清理本任务，不影响其他并发任务
                    print(f"🔌 [SSE] 客户端断开连接，清理任务 {task_id}")
                    break
        except Exception as e:
            # 🌟 异常只影响当前 task_id 的推流，不向外传播
            print(f"❌ [SSE] 任务 {task_id} 推流异常: {e}")
            yield f'data: {{"type": "error", "message": "日志推送异常: {str(e)}"}}\n\n'
        finally:
            # 🌟 无论正常结束还是异常，只清理当前任务的队列
            if task_id in task_queues:
                del task_queues[task_id]
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


# ==================== 🌟 系统配置中心 (Settings) API ====================

def _mask_key(val: Optional[str]) -> str:
    """对敏感 Key 进行脚敏处理，仅显示前 4 位和后 4 位。"""
    if not val:
        return ""
    if len(val) <= 8:
        return "****"
    return val[:4] + "****" + val[-4:]


@app.get("/api/settings")
async def get_settings():
    """返回当前系统配置（敏感 Key 已脚敏）。"""
    cfg = _config_module._cfg
    return {
        "OPENAI_API_KEY": _mask_key(cfg("OPENAI_API_KEY", "LLM_API_KEY", json_key="OPENAI_API_KEY")),
        "OPENAI_BASE_URL": cfg("OPENAI_BASE_URL", "LLM_BASE_URL", json_key="OPENAI_BASE_URL") or "",
        "SERPER_API_KEY": _mask_key(cfg("SERPER_API_KEY", json_key="SERPER_API_KEY")),
        "FEISHU_APP_ID": _mask_key(cfg("FEISHU_APP_ID", "APP_ID", json_key="FEISHU_APP_ID")),
        "FEISHU_APP_SECRET": _mask_key(cfg("FEISHU_APP_SECRET", "APP_SECRET", json_key="FEISHU_APP_SECRET")),
        "FEISHU_APP_TOKEN": _mask_key(cfg("FEISHU_APP_TOKEN", "APP_TOKEN", json_key="FEISHU_APP_TOKEN")),
    }


@app.post("/api/settings")
async def save_settings(payload: SettingsPayload):
    """保存配置到 settings.json，并热重载 config 模块及全局客户端。"""
    SETTINGS_DATA_PATH.parent.mkdir(parents=True, exist_ok=True)

    existing: dict = {}
    if SETTINGS_DATA_PATH.exists():
        try:
            existing = json.loads(SETTINGS_DATA_PATH.read_text(encoding="utf-8"))
        except Exception:
            pass

    updates = {k: v for k, v in payload.model_dump(exclude_none=True).items() if v and str(v).strip()}
    existing.update(updates)
    SETTINGS_DATA_PATH.write_text(json.dumps(existing, ensure_ascii=False, indent=2), encoding="utf-8")

    importlib.reload(_config_module)

    global client, APP_ID, APP_SECRET, APP_TOKEN, LLM_API_KEY, LLM_BASE_URL, SERPER_API_KEY
    LLM_API_KEY = _config_module.LLM_API_KEY
    LLM_BASE_URL = _config_module.LLM_BASE_URL
    APP_ID = _config_module.FEISHU_APP_ID
    APP_SECRET = _config_module.FEISHU_APP_SECRET
    APP_TOKEN = _config_module.FEISHU_APP_TOKEN
    SERPER_API_KEY = _config_module.SERPER_API_KEY
    client = OpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL) if LLM_API_KEY else None

    return {"status": "ok", "message": f"已保存 {len(updates)} 个配置项并热重载成功"}

# ==================== 🌟 实验性功能：多Agent简历协作工厂 API ====================
@app.post("/api/agents/deep-rewrite")
async def multi_agent_rewrite(payload: MultiAgentRewriteRequest):
    print(f"\n====== 📡 收到【多Agent深度改写】请求 ======")
    try:
        initial_state = {
            "original_full_text": payload.original_resume,
            "jd": payload.jd_text,
            "diagnosis_report": "测试模式无诊断报告",
            "parsed_blocks": [],
            "working_rewritten_blocks": [],
            "critic_feedback": "",
            "current_score": 0,
            "revision_count": 0,
            "token_usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0},
            "logs": [],
            "final_markdown": ""
        }
        
        result = await asyncio.to_thread(multi_agent_app.invoke, initial_state)
        
        print(f"🎉 多Agent协作完成！共经历 {result.get('revision_count', 0)} 轮重写。")
        
        return {
            "status": "success",
            "final_version": result.get("final_markdown", ""),
            "meeting_logs": result.get("logs", []),  
            "revision_rounds": result.get("revision_count", 0)
        }
        
    except Exception as e:
        print(f"❌ 多Agent改写测试接口失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"多Agent引擎报错: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True,access_log=False)